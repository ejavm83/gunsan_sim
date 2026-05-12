"""군산 공장 하이브리드 공정 시뮬레이션 웹 대시보드.

Streamlit 기반 인터랙티브 대시보드로, 브라우저에서 파라미터를 조정하고
시뮬레이션을 실행한 뒤 결과를 즉시 확인할 수 있다.

실행 방법::

    streamlit run webapp.py
    # 또는
    py -3 -m streamlit run webapp.py
"""

from __future__ import annotations

import dataclasses
import hashlib
import html as html_module
import json
import os
import re
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas  # noqa: F401 — plotly가 부분 초기화된 pandas를 보지 않도록 먼저 로드

from docx import Document

import streamlit as st
import streamlit.components.v1 as components
from extra_streamlit_components import CookieManager
import plotly.graph_objects as go
from plotly.subplots import make_subplots

from config import (
    DEFAULT_CONFIG,
    InboundConfig,
    SortingConfig,
    MeltingConfig,
    CastingConfig,
    OutboundConfig,
    SimulationConfig,
)
from metrics import Metrics
from simulation import run_simulation
from report import (
    analyze,
    Analysis,
    build_layperson_visual_figures,
    layperson_interpretation_export_markdown,
    layperson_interpretation_markdown,
)
from pdf_report import markdown_simulation_report_to_pdf

_BUILD_INFO_TEXT = "(주) 지엠티 김길용 수석, v0.0.2 (2026.05.09)"
_REPO_ROOT = Path(__file__).resolve().parent
_PROCESS_DETAIL_MD = _REPO_ROOT / "군산 공정 상세-김홍태보완.md"
_SIMPY_CPSAT_MD = _REPO_ROOT / "docs" / "simpy_cpsat_overview.md"
_TERMS_GLOSSARY_MD = _REPO_ROOT / "docs" / "terms_glossary.md"
_SIM_INPUTS_CONSTRAINTS_MD = _REPO_ROOT / "docs" / "simulation_inputs_constraints.md"

# 상단 한 줄 내비(실행·결과 + 문서). 예전 세션의 "시뮬레이션" 값은 실행·결과로 치환한다.
_MAIN_NAV_HOME = "실행·결과"
_MAIN_NAV_OPTIONS: tuple[str, ...] = (
    _MAIN_NAV_HOME,
    "📚 용어 및 약어",
    "🏭 공정 상세",
    "🔬 방법론 및 라이브러리",
    "📌 입력과 규칙 (쉬운 설명)",
)

# Plotly 호버 박스: 흰 plot 배경과 구분되도록 어두운 패널 + 밝은 테두리
_PLOTLY_HOVERLABEL = dict(
    bgcolor="rgb(15,23,42)",
    bordercolor="rgb(56,189,248)",
    font=dict(color="rgb(241,245,249)", size=14),
)


def _style_layperson_result_figure(fig: go.Figure) -> go.Figure:
    """대시보드 어두운 톤에 맞춘 일반인 해석용 Plotly 스타일."""
    fig.update_layout(
        template="plotly_dark",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(15,23,42,0.72)",
        font=dict(color="#e2e8f0", size=12),
        title_font=dict(color="#f8fafc", size=14),
        legend=dict(font=dict(color="#cbd5e1")),
        hoverlabel=_PLOTLY_HOVERLABEL,
    )
    fig.update_xaxes(
        gridcolor="rgba(148,163,184,0.22)",
        zerolinecolor="rgba(148,163,184,0.35)",
        color="#94a3b8",
    )
    fig.update_yaxes(
        gridcolor="rgba(148,163,184,0.22)",
        zerolinecolor="rgba(148,163,184,0.35)",
        color="#94a3b8",
    )
    return fig

# Streamlit 위젯 help= 트리거(ⓘ)와 호버 패널 — 기본은 대비가 약해 안내 버튼·툴팁으로 보이게 조정
_STREAMLIT_HELP_TOOLTIP_CSS = """
<style>
[data-testid="stTooltipIcon"] button {
  cursor: help !important;
  border-radius: 999px !important;
  border: 1.5px solid rgba(34, 211, 238, 0.45) !important;
  background: rgba(34, 211, 238, 0.12) !important;
  color: rgb(165, 243, 252) !important;
  min-width: 1.7rem !important;
  min-height: 1.7rem !important;
  transition: background 0.15s ease, border-color 0.15s ease, transform 0.12s ease,
    box-shadow 0.12s ease !important;
}
[data-testid="stTooltipIcon"] button:hover,
[data-testid="stTooltipIcon"] button:focus-visible {
  background: rgba(34, 211, 238, 0.22) !important;
  border-color: rgb(34, 211, 238) !important;
  box-shadow: 0 0 0 2px rgba(34, 211, 238, 0.22) !important;
  transform: scale(1.06) !important;
  outline: none !important;
}
[data-testid="stTooltipIcon"] button:focus-visible {
  box-shadow: 0 0 0 3px rgba(34, 211, 238, 0.35) !important;
}
[data-testid="stTooltipIcon"] svg.icon {
  width: 1.15em !important;
  height: 1.15em !important;
}
[data-testid="stTooltipContent"],
[data-testid="stTooltipErrorContent"] {
  max-width: min(28rem, 92vw) !important;
  padding: 0.85rem 1rem !important;
  font-size: 0.9375rem !important;
  line-height: 1.55 !important;
  border-radius: 0.625rem !important;
  user-select: text !important;
  -webkit-user-select: text !important;
  -moz-user-select: text !important;
  cursor: auto !important;
}
/* 본문 slate-900(15,23,42)과 동일 톤이면 패널이 묻혀 보이므로 한 단계 올린 카드 느낌으로 분리 */
[data-testid="stTooltipContent"] {
  background: linear-gradient(
    165deg,
    rgb(71, 85, 105) 0%,
    rgb(51, 65, 85) 42%,
    rgb(38, 49, 68) 100%
  ) !important;
  border: 2px solid rgba(34, 211, 238, 0.65) !important;
  color: rgb(248, 250, 252) !important;
  box-shadow:
    0 0 0 1px rgba(15, 23, 42, 0.9),
    0 0 0 5px rgba(34, 211, 238, 0.14),
    0 14px 36px rgba(2, 6, 23, 0.65),
    0 4px 12px rgba(2, 6, 23, 0.35) !important;
}
[data-testid="stTooltipErrorContent"] {
  background: linear-gradient(
    165deg,
    rgb(127, 29, 29) 0%,
    rgb(91, 20, 20) 50%,
    rgb(69, 10, 10) 100%
  ) !important;
  border: 2px solid rgba(252, 165, 165, 0.85) !important;
  color: rgb(254, 242, 242) !important;
  box-shadow:
    0 0 0 1px rgba(69, 10, 10, 0.95),
    0 0 0 5px rgba(248, 113, 113, 0.18),
    0 14px 36px rgba(2, 6, 23, 0.55),
    0 4px 12px rgba(69, 10, 10, 0.45) !important;
}
[data-testid="stTooltipContent"] *,
[data-testid="stTooltipErrorContent"] * {
  user-select: text !important;
  -webkit-user-select: text !important;
  -moz-user-select: text !important;
}
/* Plotly 호버 패널: 드래그로 텍스트 선택 후 복사 가능하도록 */
[data-testid="stPlotlyChart"] .hovertext,
[data-testid="stPlotlyChart"] .hovertext * {
  user-select: text !important;
  -webkit-user-select: text !important;
  -moz-user-select: text !important;
  cursor: auto !important;
}
</style>
"""

# 표·슬라이더·사이드바 리듬 등(툴팁 표시 여부와 무관하게 항상 적용)
_DASHBOARD_BASE_CSS = """
<style>
[data-testid="stMarkdownContainer"] table {
  border-collapse: collapse !important;
  width: 100% !important;
  margin: 0.35rem 0 0.85rem 0 !important;
  font-size: 0.94rem !important;
}
[data-testid="stMarkdownContainer"] thead th {
  background: rgba(30, 41, 59, 0.9) !important;
  color: rgb(226, 232, 240) !important;
  font-weight: 600 !important;
}
[data-testid="stMarkdownContainer"] th,
[data-testid="stMarkdownContainer"] td {
  border: 1px solid rgba(100, 116, 139, 0.5) !important;
  padding: 0.5rem 0.65rem !important;
  vertical-align: top !important;
}
[data-testid="stMarkdownContainer"] tbody tr:nth-child(even) td {
  background: rgba(15, 23, 42, 0.4) !important;
}
[data-testid="stSlider"] [data-baseweb="slider"] [role="slider"] {
  background-color: rgb(34, 211, 238) !important;
  border: 2px solid rgb(8, 145, 178) !important;
  box-shadow: 0 0 0 1px rgba(15, 23, 42, 0.35) !important;
}
section[data-testid="stSidebar"] [data-testid="stHeader"] {
  margin-top: 0.2rem !important;
  margin-bottom: 0.15rem !important;
}
section[data-testid="stSidebar"] > div > div > div > div.block-container {
  padding-top: 0.45rem !important;
  padding-bottom: 0.6rem !important;
  padding-left: 0.7rem !important;
  padding-right: 0.7rem !important;
}
section[data-testid="stSidebar"] [data-testid="stCaption"] {
  line-height: 1.55 !important;
  margin-top: 0.2rem !important;
  margin-bottom: 0.45rem !important;
}
section[data-testid="stSidebar"] [data-testid="stHeader"] + div [data-testid="stCaption"] {
  margin-top: 0.15rem !important;
}
div[data-testid="stAlert"] {
  border-radius: 0.5rem !important;
}
.gunsan-callout {
  border-radius: 0.625rem;
  padding: 0.95rem 1.1rem;
  border: 1px solid rgba(34, 211, 238, 0.35);
  background: linear-gradient(135deg, rgba(8, 47, 73, 0.55) 0%, rgba(15, 23, 42, 0.72) 100%);
  color: rgb(226, 232, 240);
  font-size: 0.97rem;
  line-height: 1.58;
  margin: 0.35rem 0 0.9rem 0;
}
.gunsan-callout strong {
  color: rgb(240, 249, 255);
}
/* 세부공정 프로세스: 제목·설명·파이프라인 직전까지 세로 여백 축소 */
div[data-testid="stMarkdownContainer"] .gunsan-pf-section-wrap {
  margin: -0.2rem 0 0;
}
div[data-testid="stMarkdownContainer"] .gunsan-pf-section-wrap h3 {
  margin: 0 0 0.18rem !important;
  padding: 0 !important;
  font-size: 1.32rem;
  font-weight: 600;
  line-height: 1.28;
  color: rgb(248, 250, 252);
}
div[data-testid="stMarkdownContainer"] .gunsan-pf-section-wrap .gunsan-pf-section-cap {
  margin: 0 0 0.3rem !important;
  padding: 0 !important;
  font-size: 0.82rem;
  line-height: 1.45;
  color: rgb(148, 163, 184);
}
div[data-testid="stElementContainer"]:has(.gunsan-pf-section-wrap)
  + div[data-testid="stElementContainer"] {
  margin-top: -0.4rem !important;
}
/* 메인 본문: 가로 공간 최대 활용(와이드 모드 보조) + 기본 블록 여백 완화 */
section[data-testid="stMain"] > div {
  max-width: 100% !important;
}
section[data-testid="stMain"] .block-container {
  max-width: 100% !important;
  padding-top: 0.4rem !important;
  padding-bottom: 0.55rem !important;
  padding-left: 0.65rem !important;
  padding-right: 0.65rem !important;
}
section[data-testid="stMain"] {
  width: 100% !important;
}
/* components.html: 기본 래퍼가 iframe보다 넓은 높이를 잡지 않도록 여백만 완화 */
div[data-testid="stIFrame"] {
  margin-bottom: 0.15rem !important;
}
/* 메인 본문 마크다운: 장문 읽기 피로 완화(살짝 낮은 대비 + 넉넉한 행간·목록 간격) */
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] p {
  color: rgb(203, 213, 225) !important;
  line-height: 1.72 !important;
  margin: 0.4em 0 0.7em 0 !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] li > p {
  margin: 0.28em 0 !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] h1,
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] h2 {
  color: rgb(248, 250, 252) !important;
  line-height: 1.32 !important;
  margin-top: 1.15rem !important;
  margin-bottom: 0.45rem !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] h3 {
  color: rgb(241, 245, 249) !important;
  line-height: 1.38 !important;
  margin-top: 0.52rem !important;
  margin-bottom: 0.38rem !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] ol,
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] ul {
  margin: 0.45rem 0 1rem 0 !important;
  padding-left: 1.4rem !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] li {
  margin: 0.55rem 0 0.9rem 0 !important;
  line-height: 1.68 !important;
  color: rgb(203, 213, 225) !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] li li {
  margin: 0.28rem 0 0.4rem 0 !important;
}
section[data-testid="stMain"] [data-testid="stMarkdownContainer"] li strong {
  color: rgb(240, 249, 255) !important;
  font-weight: 600 !important;
}
section[data-testid="stSidebar"] [data-testid="stSlider"] {
  margin-bottom: 0.75rem !important;
  padding-bottom: 0.15rem !important;
}
section[data-testid="stSidebar"] [data-testid="stSlider"] label {
  margin-bottom: 0.35rem !important;
}
</style>
"""


def _inject_dashboard_base_styles() -> None:
    st.markdown(_DASHBOARD_BASE_CSS, unsafe_allow_html=True)


def _inject_streamlit_help_tooltip_styles(*, show_widget_help: bool) -> None:
    if show_widget_help:
        st.markdown(_STREAMLIT_HELP_TOOLTIP_CSS, unsafe_allow_html=True)
    else:
        st.markdown(
            '<style>[data-testid="stTooltipIcon"]{display:none!important;}</style>',
            unsafe_allow_html=True,
        )


# 마지막 시뮬 결과 보존(Streamlit 버튼은 rerun당 한 번만 True)
_GUNSAN_LAST_RUN_BUNDLE_KEY = "gunsan_last_run_bundle"


def _read_markdown_file(path: Path, default_text: str = "") -> str:
    if not path.is_file():
        return default_text
    return path.read_text(encoding="utf-8")


def _save_markdown_file(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    # 서버 파일에 저장되므로 같은 서버를 보는 다른 사용자도 동일 내용을 확인할 수 있다.
    path.write_text(content, encoding="utf-8")


def _render_mermaid(chart: str, *, height: int = 420) -> None:
    """호환성 이슈가 있어 Mermaid 원문은 접어두고 Graphviz로 렌더한다."""
    _ = chart
    _ = height


def _render_graphviz_chart(
    dot_source: str,
    *,
    height: int = 480,
    fit: bool = True,
    scrolling: bool = True,
) -> None:
    """브라우저 콘솔 worker 경고를 피하기 위해 d3-graphviz를 직접 렌더링한다.

    d3-graphviz는 ``fit`` 이 참이어도 SVG ``width``/``height`` 가 없으면 스케일이 적용되지 않아
    기본(pt) 크기로 그려져 iframe 안에서 잘릴 수 있다. 컨테이너 픽셀 크기를 넘겨 전체가 보이게 한다.
    """
    container_id = f"gv-{hashlib.md5(dot_source.encode('utf-8')).hexdigest()[:10]}"
    dot_json = json.dumps(dot_source)
    fit_js = "true" if fit else "false"
    html = f"""
<div id="{container_id}" style="width:100%;height:{height}px;box-sizing:border-box;"></div>
<script src="https://unpkg.com/d3@7/dist/d3.min.js"></script>
<script src="https://unpkg.com/@hpcc-js/wasm@2.20.0/dist/graphviz.umd.js"></script>
<script src="https://unpkg.com/d3-graphviz@5/build/d3-graphviz.js"></script>
<script>
  const dot = {dot_json};
  const el = document.getElementById("{container_id}");
  const w = Math.max(el.clientWidth, 320);
  const h = Math.max(el.clientHeight, 200);
  d3.select("#{container_id}")
    .graphviz({{ useWorker: false, fit: {fit_js}, zoom: false, width: w, height: h }})
    .renderDot(dot);
</script>
"""
    components.html(html, height=height, scrolling=scrolling)


def _format_kpi_minutes(minutes: float) -> str:
    """KPI용 분 표기. 2시간 미만은 소수 첫째 자리 분, 그 이상은 일·시간·분(앞쪽 0 단위 생략)."""
    if minutes < 0:
        minutes = 0.0
    if minutes < 120:
        return f"{minutes:.1f} 분"
    total_sec = int(round(minutes * 60))
    d, rem = divmod(total_sec, 86400)
    h, rem2 = divmod(rem, 3600)
    m, _ = divmod(rem2, 60)
    parts: list[str] = []
    if d:
        parts.append(f"{d}일")
    if h or d:
        parts.append(f"{h}시간")
    parts.append(f"{m}분")
    return " ".join(parts)


def _sim_clock_label(time_min: float) -> str:
    """시뮬레이션 0분 기준 경과 시각(일차 + 시:분)."""
    if time_min < 0:
        time_min = 0.0
    day = int(time_min // (24 * 60)) + 1
    rem = time_min % (24 * 60)
    total_m = int(round(rem))
    total_m = max(0, min(total_m, 24 * 60 - 1))
    h, m = divmod(total_m, 60)
    return f"{day}일차 {h:02d}:{m:02d}"


# 영문 stage/kind 를 일반인이 바로 알아볼 수 있는 한글로 옮기는 표.
# 시뮬 코드(`simulation.py`)는 영문 키를 그대로 쓰고, 표시 단계에서만 한글로 바꾼다.
_EVENT_STAGE_KOR: dict[str, str] = {
    "inbound": "① 원료 입고",
    "sorting": "② 선별",
    "press": "② 압착·파레트",
    "melting": "③ 용해(반사로)",
    "casting": "④ 주조",
    "outbound": "⑤ 제품 출하",
}

_EVENT_KIND_KOR: dict[tuple[str, str], str] = {
    ("inbound", "arrive"): "입고 트럭 도착",
    ("inbound", "weigh_in"): "1차 계근(공차)",
    ("inbound", "unloaded"): "하역 완료",
    ("inbound", "depart"): "입고 트럭 출차",
    ("sorting", "sort_done"): "선별 완료(8개 더미)",
    ("press", "pallet_done"): "파레트 1개 생산",
    ("melting", "batch_collected"): "1배치 분량 모음(32 파레트)",
    ("melting", "elevator_done"): "엘리베이터 장입 완료",
    ("melting", "melt_start"): "반사로 용해 시작",
    ("melting", "melt_done"): "반사로 용해 완료",
    ("melting", "batch_done"): "1배치 완료",
    ("casting", "flake_buffer_full"): "후레이크 버퍼 가득 참(생산 정지)",
    ("casting", "flake_done"): "후레이크 주조 완료",
    ("casting", "scr_buffer_full"): "SCR 코일 버퍼 가득 참(생산 정지)",
    ("casting", "scr_done"): "SCR 코일 주조 완료",
    ("outbound", "arrive"): "출하 트럭 도착",
    ("outbound", "depart"): "출하 트럭 출차",
}

_EVENT_LOAD_KIND_KOR: dict[str, str] = {
    "flake": "후레이크(1 t 포대)",
    "scr": "SCR 코일(4 t/개)",
}


def _humanize_event_detail(stage: str, kind: str, detail: dict) -> str:
    """이벤트 `detail` 딕셔너리를 사람이 읽기 쉬운 한 줄 한글 설명으로 만든다."""
    if not detail:
        return ""
    parts: list[str] = []
    if "truck" in detail:
        parts.append(f"{detail['truck']}번 트럭")
    if "furnace" in detail:
        parts.append(f"{detail['furnace']}호 반사로")
    if "pallet" in detail:
        parts.append(f"{detail['pallet']}번 파레트")
    if "pallets" in detail:
        parts.append(f"파레트 {detail['pallets']}개")
    if "trips" in detail:
        parts.append(f"엘리베이터 {detail['trips']}회 운반")
    if "units" in detail:
        if kind == "flake_done":
            parts.append(f"후레이크 {detail['units']}포대")
        elif kind == "scr_done":
            parts.append(f"SCR 코일 {detail['units']}개")
        else:
            parts.append(f"{detail['units']}개")
    if "load_kind" in detail:
        parts.append(_EVENT_LOAD_KIND_KOR.get(str(detail["load_kind"]), str(detail["load_kind"])))
    if "ton" in detail:
        try:
            parts.append(f"적재 {float(detail['ton']):.2f} t")
        except (TypeError, ValueError):
            parts.append(f"적재 {detail['ton']} t")
    if "buffer" in detail:
        parts.append(f"버퍼 잔량 {detail['buffer']}")
    # 위 매핑에 없는 키는 끝에 `키=값` 으로 덧붙여 정보 손실을 피한다.
    known = {"truck", "furnace", "pallet", "pallets", "trips", "units",
             "load_kind", "ton", "buffer"}
    extras = [f"{k}={v}" for k, v in detail.items() if k not in known]
    if extras:
        parts.append("(" + ", ".join(extras) + ")")
    return ", ".join(parts)


def _build_events_log_df(metrics: Metrics) -> "pandas.DataFrame":
    """`metrics.events` 를 표·다운로드용 테이블로 변환한다.

    화면에는 영문 stage/kind 대신 일반인이 바로 이해할 수 있는 한글 명칭과
    풀어 쓴 상세 설명을 보여 준다. 분석·CSV 호환을 위해 원본 영문 값도 함께
    `구간(원문)`·`사건(원문)` 컬럼으로 보존한다.
    """
    rows: list[dict[str, object]] = []
    for i, ev in enumerate(metrics.events, start=1):
        stage_kor = _EVENT_STAGE_KOR.get(ev.stage, ev.stage)
        kind_kor = _EVENT_KIND_KOR.get((ev.stage, ev.kind), ev.kind)
        rows.append(
            {
                "#": i,
                "시각": _sim_clock_label(ev.time_min),
                "경과(분)": round(ev.time_min, 1),
                "구간": stage_kor,
                "사건": kind_kor,
                "상세 설명": _humanize_event_detail(ev.stage, ev.kind, ev.detail or {}),
                "구간(원문)": ev.stage,
                "사건(원문)": ev.kind,
            }
        )
    return pandas.DataFrame(rows)


def _markdown_escape_cell(value: object) -> str:
    """Markdown 표 셀: 줄바꿈·세로줄을 깨지 않게 정리한다."""
    return str(value).replace("\r\n", " ").replace("\n", " ").replace("|", "\\|")


def _dataframe_to_markdown_table(
    df: "pandas.DataFrame",
    *,
    max_rows: int | None = None,
) -> str:
    """pandas `to_markdown` 없이 GFM 스타일 표를 만든다."""
    if df is None or df.empty:
        return "_표시할 행이 없습니다._\n"
    view = df if max_rows is None else df.iloc[: max_rows if max_rows > 0 else 0]
    if view.empty:
        return "_표시할 행이 없습니다._\n"
    cols = list(view.columns)
    header = "| " + " | ".join(_markdown_escape_cell(c) for c in cols) + " |"
    sep = "| " + " | ".join("---" for _ in cols) + " |"
    body_lines: list[str] = []
    for _, row in view.iterrows():
        body_lines.append(
            "| " + " | ".join(_markdown_escape_cell(row[c]) for c in cols) + " |"
        )
    out = "\n".join([header, sep, *body_lines]) + "\n"
    if max_rows is not None and len(df) > max_rows:
        out += f"\n_(최대 {max_rows:,}행만 포함. 전체 {len(df):,}행.)_\n"
    return out


def _build_simulation_results_markdown(
    metrics: Metrics,
    cfg: SimulationConfig,
    analysis: Analysis,
    ev_df_view: "pandas.DataFrame",
    ev_df_total_rows: int,
    event_filter_note: str,
    kpi_specs: list[tuple[str, str, str | None]],
    *,
    generated_at: str,
) -> str:
    """「마지막 실행 결과」 화면에 대응하는 텍스트 보고서(차트 제외)."""
    s = analysis.summary
    flake_pct = int(round(cfg.casting.flake_ratio * 100))
    cfg_rows = [
        ("시뮬레이션 일수", str(cfg.sim_days), "일"),
        ("랜덤 시드", str(cfg.random_seed), ""),
        ("일 트럭 수", str(cfg.inbound.trucks_per_day), "대"),
        ("트럭 적재", str(cfg.inbound.payload_ton), "t"),
        ("하역 베이", str(cfg.inbound.unloading_bays), ""),
        ("선별 워커", str(cfg.sorting.sorters), ""),
        ("압착기", str(cfg.sorting.press_machines), ""),
        ("파레트 버퍼", str(cfg.sorting.pallet_buffer_capacity), "개"),
        ("반사로", str(cfg.melting.furnace_count), "대"),
        ("배치 단위", str(cfg.melting.batch_ton), "t"),
        ("큐프레이크 비율", f"{flake_pct}%", ""),
        ("출하 트럭 평균 간격", str(int(round(cfg.outbound.empty_truck_interval_min))), "분"),
    ]
    cfg_table = (
        "| 파라미터 | 값 | 단위 |\n|---|---|---|\n"
        + "\n".join(
            f"| {_markdown_escape_cell(n)} | {_markdown_escape_cell(v)} | {_markdown_escape_cell(u)} |"
            for n, v, u in cfg_rows
        )
    )
    kpi_table = (
        "| 지표 | 값 | 참고(웹 도움말 요약) |\n|---|---|---|\n"
        + "\n".join(
            f"| {_markdown_escape_cell(lbl)} | {_markdown_escape_cell(val)} | "
            f"{_markdown_escape_cell(h or '')} |"
            for lbl, val, h in kpi_specs
        )
    )
    util_table = (
        "| 자원 | 가동률(%) |\n|---|---|\n"
        + "\n".join(
            f"| {_markdown_escape_cell(name)} | {pct * 100:.1f} |"
            for name, pct in sorted(analysis.util.items(), key=lambda kv: kv[1], reverse=True)
        )
    )
    buf_data = []
    for name, stats in analysis.queue_stats.items():
        buf_data.append(
            {
                "버퍼": name,
                "평균 점유": f"{stats['avg']:.1f}",
                "최대 점유": f"{stats['max']:.0f}",
                "95퍼센타일": f"{stats['p95']:.0f}",
            }
        )
    buf_md = _dataframe_to_markdown_table(pandas.DataFrame(buf_data))

    ev_cols = [c for c in ev_df_view.columns if c in ("#", "시각", "경과(분)", "구간", "사건", "상세 설명", "구간(원문)", "사건(원문)")]
    if not ev_cols and not ev_df_view.empty:
        ev_cols = list(ev_df_view.columns)
    ev_slice = ev_df_view[ev_cols] if ev_cols else ev_df_view
    ev_cap = 5000
    ev_md = _dataframe_to_markdown_table(ev_slice, max_rows=ev_cap)

    daily_lines: list[str] = []
    if analysis.daily_throughput_ton:
        daily_lines.append("### 일별 생산량 (톤)\n")
        daily_lines.append("| 일차 | 큐프레이크(t) | SCR(t) |\n|---|---|---|")
        for d, fl, sc in analysis.daily_throughput_ton:
            daily_lines.append(f"| {d} | {fl:.1f} | {sc:.1f} |")
        daily_lines.append("")

    insights_md = "\n".join(f"- {line}" for line in analysis.insights) if analysis.insights else "_없음_"
    rec_md = "\n".join(f"- {line}" for line in analysis.recommendations) if analysis.recommendations else "_없음_"
    truck_md = "\n".join(f"- {line}" for line in analysis.truck_flow_insights) if analysis.truck_flow_insights else "_없음_"

    parts: list[str] = [
        "# 군산 SCR 공정 물류 시뮬레이션 — 실행 보고서",
        "",
        f"- 생성 시각: `{generated_at}`",
        f"- 빌드: {_BUILD_INFO_TEXT}",
        "",
        "## 1. 실행 조건",
        "",
        cfg_table,
        "",
        "## 2. 핵심 KPI",
        "",
        kpi_table,
        "",
        "## 3. 일반인을 위한 상세 결과 해석",
        "",
        layperson_interpretation_export_markdown(metrics, cfg, analysis),
        "",
        "## 4. 병목 진단",
        "",
        f"- **식별된 병목:** {_markdown_escape_cell(analysis.bottleneck)}",
        f"- **근거:** {_markdown_escape_cell(analysis.bottleneck_reason)}",
        "",
        "## 5. 자원 가동률",
        "",
        util_table,
        "",
        "## 6. 버퍼 통계",
        "",
        buf_md,
        "",
        *daily_lines,
        "## 7. 시뮬레이션 분석 결과 인사이트",
        "",
        "### 관찰 포인트",
        "",
        insights_md,
        "",
        "### 권장 액션",
        "",
        rec_md,
        "",
        "### 트럭·물류 시사점",
        "",
        truck_md,
        "",
        "## 8. 시간순 사건 로그",
        "",
        f"- 화면과 동일 필터 요약: {event_filter_note}",
        f"- 표시 행 수: **{len(ev_df_view):,}** / 전체 **{ev_df_total_rows:,}**",
        "",
        ev_md,
        "",
        "## 9. 차트에 대해",
        "",
        "Plotly 차트(가동률, 버퍼 시계열, 반사로 Gantt, 누적 트럭, 체류 분포, 일별 생산 등)는 "
        "벡터 그래픽이라 이 Markdown에는 넣지 않았습니다. 웹 대시보드 각 차트 아래 "
        "「HTML 다운로드」로 동일 실행 결과의 그래프를 저장할 수 있습니다.",
        "",
    ]
    return "\n".join(parts).strip() + "\n"


def _render_simpy_cpsat_overview_for_dashboard() -> None:
    """docs/simpy_cpsat_overview.md 와 동일 요지를 웹 탭에 표시한다."""
    st.markdown("---")
    st.markdown("### SimPy · CP-SAT — 역할, 관계, 활용 시점")
    st.markdown("""
이 프로젝트는 공장 전체 시간 흐름을 SimPy로 재현하고, 터미널(`python main.py`)에서만
선택적으로 반사로 배치 스케줄을 CP-SAT로 최적화해 이론적 메이크스팬과 비교합니다.
두 도구는 역할이 다르고, 이 대시보드 버튼으로 돌리는 시뮬은 SimPy만 해당합니다.
    """)

    _role_compare = pandas.DataFrame(
        [
            {
                "구분": "본질",
                "SimPy": "이산사건 시뮬: 자원·버퍼·확률 출하까지 시간 순서 재현",
                "CP-SAT (저장소에서의 역할)": (
                    "제약+최적화: 배치를 두 반사로에 배정해 전체 완료 시각을 최소화하는 정수계획"
                ),
            },
            {
                "구분": "여기서 얻는 것",
                "SimPy": "이벤트·KPI·버퍼 동역학",
                "CP-SAT (저장소에서의 역할)": "배치 시작 시각표·최소 메이크스팬(수학 모델 해)",
            },
            {
                "구분": "반사로 2대",
                "SimPy": "SimPy Resource 선착순 근사",
                "CP-SAT (저장소에서의 역할)": "반사로별 겹침 없이 시작 시각 최적화",
            },
            {
                "구분": "실행 위치",
                "SimPy": "run_simulation() — 웹·CLI 공통",
                "CP-SAT (저장소에서의 역할)": (
                    "main.py만 (기본 켜짐, --no-optimize로 끔). 웹 실행 버튼과는 연동되지 않음"
                ),
            },
            {
                "구분": "코드",
                "SimPy": "simulation.py, metrics.py",
                "CP-SAT (저장소에서의 역할)": "optimizer.py, main.py",
            },
        ]
    )
    with st.container(border=False):
        st.caption("SimPy vs CP-SAT 한눈에 비교 (행 선택·복사 가능)")
        st.dataframe(
            _role_compare,
            use_container_width=True,
            hide_index=True,
            height=min(260, 38 + len(_role_compare) * 36),
        )

    st.markdown(
        '<div class="gunsan-callout">'
        "<strong>이 화면의 KPI·차트</strong>는 모두 SimPy 결과입니다. "
        "CP-SAT 스케줄·메이크스팬 비교는 저장소에서 <code>python main.py</code>로 실행하세요 "
        "(웹에 붙이려면 별도 연동 개발이 필요합니다)."
        "</div>",
        unsafe_allow_html=True,
    )
    if _SIMPY_CPSAT_MD.is_file():
        st.caption(f"상세 전문: 저장소 `{_SIMPY_CPSAT_MD.relative_to(_REPO_ROOT)}`")

    with st.expander("그림 1 — SimPy(항상)와 CP-SAT(CLI 선택) 역할 분담", expanded=False):
        st.caption(
            "펼치면 큰 캔버스에서 전체 흐름을 봅니다. 왼쪽: 시간축 시뮬 / 오른쪽: 배치 단위 반사로 재스케줄"
        )
        _render_graphviz_chart(
            r"""
digraph G {
  rankdir=LR;
  graph [pad=0.2, nodesep=0.4, ranksep=0.6];
  node [shape=box, style="rounded,filled", fillcolor="#f8fafc", color="#475569", fontname="Malgun Gothic"];

  Cfg [label="설정 입력\nSimulationConfig\n(config.py)"];
  S1 [label="GunsanFactory\n자원·Store·프로세스", fillcolor="#e8f5e9"];
  S2 [label="env.run()\n버퍼·출하 간격 포함", fillcolor="#e8f5e9"];
  S3 [label="Metrics\n이벤트·대기·처리량", fillcolor="#e8f5e9"];
  R1 [label="배치 시작 후보 시각\n(이벤트 또는 추정)", fillcolor="#fff3e0"];
  R2 [label="estimate_batch_duration\n배치 1건 길이(분)", fillcolor="#fff3e0"];
  R3 [label="solve_furnace_schedule\n메이크스팬 최소화", fillcolor="#fff3e0"];
  R4 [label="콘솔 스케줄·메이크스팬", fillcolor="#fff3e0"];

  Cfg -> S1;
  S1 -> S2 -> S3;
  S3 -> R1 [label="press/pallet_done 실측"];
  Cfg -> R1 [label="실측 없으면 근사"];
  Cfg -> R2;
  R1 -> R3;
  R2 -> R3;
  R3 -> R4;
}
""",
            height=620,
            scrolling=True,
        )

    with st.expander("그림 2 — `python main.py` 실행 순서(요약)", expanded=False):
        _render_graphviz_chart(
            r"""
digraph G {
  rankdir=TB;
  graph [pad=0.2, nodesep=0.35, ranksep=0.45];
  node [shape=box, style="rounded,filled", fillcolor="#f8fafc", color="#475569", fontname="Malgun Gothic"];

  Start [shape=oval, label="python main.py"];
  Sim [label="SimPy 실행\nrun_simulation"];
  KPI [label="요약 출력\nmetrics.summary"];
  Branch [shape=diamond, label="--no-optimize ?"];
  CP [label="CP-SAT\nestimate + solve"];
  Skip [label="CP 단계 생략"];
  Compare [label="시뮬 batch_done 최대 시각과\n메이크스팬 비교 출력"];
  Post [shape=diamond, label="후속 옵션"];
  E [label="--events CSV"];
  P [label="matplotlib 차트"];
  A [label="--animate"];
  R [label="--report HTML"];
  End [shape=oval, label="종료"];

  Start -> Sim -> KPI -> Branch;
  Branch -> CP [label="아니오"];
  Branch -> Skip [label="예"];
  CP -> Compare -> Post;
  Skip -> Post;
  Post -> E;
  Post -> P;
  Post -> A;
  Post -> R;
  E -> End;
  P -> End;
  A -> End;
  R -> End;
}
""",
            height=520,
            scrolling=True,
        )

    with st.expander("그림 3 — 직관 비유 (도로 전체 vs 관제 순서만)", expanded=False):
        _render_graphviz_chart(
            r"""
digraph G {
  rankdir=LR;
  graph [pad=0.2, nodesep=0.4, ranksep=0.6];
  node [shape=box, style="rounded,filled", fillcolor="#f8fafc", color="#475569", fontname="Malgun Gothic"];

  SIM [label="SimPy\n신호등·차선·합류\n무작위 교통량까지", fillcolor="#e8f5e9"];
  OPT [label="CP-SAT\n관제에서\n간선 순서만 최적화", fillcolor="#fff3e0"];
  DIFF [shape=oval, label="다름", fillcolor="#fee2e2", color="#b91c1c"];

  SIM -> DIFF [style=dashed, label="같은 노선이라도\n전체 재현은 다름"];
  OPT -> DIFF [style=dashed, label="배치=노선\n반사로=차선 2개"];
}
""",
            height=420,
            scrolling=True,
        )

    with st.expander("CP-SAT — 터미널에서 하는 일 / 시뮬과의 차이", expanded=False):
        st.markdown("""
- 언제: `python main.py` 이며 `--no-optimize` 가 없을 때 `optimizer` 가 호출됩니다.
- 입력: 배치당 시작 가능 시각(release) 은 시뮬 이벤트(`press`/`pallet_done`) 우선, 없으면 추정식(`estimate_batch_releases`) — 추정은 대기열·버퍼를 단순화해 시뮬과 어긋날 수 있음.
- 주의: 솔버가 내는 시작 시각표는 SimPy를 다시 돌린 결과가 아니라, “반사로 두 대만 최적으로 돌린다면”에 대한 수학 모델의 해입니다. 격차는 선착순 근사 vs 순서 최적의 차이로 읽으면 됩니다.
        """)


def _paragraph_add_markdown_bold(paragraph, text: str) -> None:
    """문단에 별표 두 개로 감싼 굵게 구간을 반영해 run을 추가한다."""
    dbl = "*" * 2
    token = re.escape(dbl)
    parts = re.split(rf"({token}[^*]+?{token})", text)
    for part in parts:
        if part.startswith(dbl) and part.endswith(dbl) and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _markdown_lines_to_docx(
    md_text: str,
    *,
    title: str = "문서",
    source_label: str | None = None,
) -> bytes:
    """마크다운 텍스트를 단순 규칙으로 Word(.docx) 바이너리로 변환한다.

    `#`, `##`, `###` 헤딩, ``-``/``*`` 글머리, 굵게(별표 두 개) 마크업 정도만 지원한다.
    표·이미지 등은 평문으로 들어가지만, 본 프로젝트 문서는 대부분 이 규칙으로 충분히 가독성을 유지한다.
    """
    doc = Document()
    doc.add_heading(title, 0)
    if source_label:
        doc.add_paragraph(f"출처: 저장소 `{source_label}`")
    for raw in md_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _paragraph_add_markdown_bold(p, stripped[2:].strip())
        else:
            p = doc.add_paragraph()
            _paragraph_add_markdown_bold(p, stripped)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@st.cache_data(show_spinner=False)
def _markdown_to_docx_bytes_cached(
    rel_path_str: str,
    title: str,
    source_label: str,
    source_mtime_ns: int,
) -> bytes:
    """저장소 상대 경로 마크다운 파일을 .docx 바이트로 변환(파일 변경 시 자동 무효화)."""
    _ = source_mtime_ns  # 캐시 키 용도
    md = (_REPO_ROOT / rel_path_str).read_text(encoding="utf-8")
    return _markdown_lines_to_docx(md, title=title, source_label=source_label)


def _process_detail_docx_download() -> tuple[bytes, str] | None:
    if not _PROCESS_DETAIL_MD.is_file():
        return None
    return (
        _markdown_to_docx_bytes_cached(
            rel_path_str=str(_PROCESS_DETAIL_MD.relative_to(_REPO_ROOT)).replace("\\", "/"),
            title="군산 공장 하이브리드 공정 상세",
            source_label=str(_PROCESS_DETAIL_MD.relative_to(_REPO_ROOT)).replace("\\", "/"),
            source_mtime_ns=_PROCESS_DETAIL_MD.stat().st_mtime_ns,
        ),
        "군산_공정_상세.docx",
    )


def _render_simulation_inputs_layperson_guide() -> None:
    """docs/simulation_inputs_constraints.md 요지를 표·전문용어 최소화로 풀어 쓴다."""
    st.markdown(
        """
이 시뮬레이션이 쓰는 단위는 코드와 같습니다. 시간은 분, 무게는 톤(t) 입니다.
아래는 “숫자가 어디서 오는지”, “이 화면에서만 무엇을 바꾸는지”, “프로그램이 어떤 규칙으로 움직이는지”를
현장·기획 담당자도 읽기 쉽게 정리한 것입니다.
        """.strip()
    )

    st.markdown("##### 1) 숫자는 어디서 정해지나요?")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown(
            """
① 공장 기본 설계표 (`config.py`)  
입고·하역·선별·압착·용해·주조·출하의 기본 시간·대수·용량이 여기 한곳에 모여 있습니다.
공정 설명 문서의 1~5단계 수치를 데이터로 옮긴 것이라고 보면 됩니다.

② 이 웹 화면(사이드바)  
실행할 때마다, 아래 「3) 이 화면에서만」에 적힌 항목만 이 기본값 위에 덮어씁니다.
나머지는 전부 기본 설계표를 그대로 따릅니다.
            """.strip()
        )
    with c2:
        st.markdown(
            """
③ 터미널에서 `python main.py`로 돌릴 때  
기간(`--days`)과 난수(`--seed`)만 간단히 바꿀 수 있고, 나머지도 기본 설계표를 따릅니다.

④ (선택) 반사로 일정만 따로 계산할 때 — CP-SAT  
시뮬 결과 로그에서 “배치가 언제쯤 준비됐는지”를 모아 반사로에 넣는 순서만 수학적으로 맞춰 보는 보조 기능입니다.
이 웹의 실행 버튼과는 연결되어 있지 않습니다. 터미널 실행 흐름에만 들어갑니다.
            """.strip()
        )

    st.markdown("##### 2) 시뮬이 지키는 ‘공장 안의 규칙’ (쉬운 말)")
    st.markdown(
        """
- 계근대는 하나로 쓰는 도로처럼 생각하세요. 입고 트럭과 출하 트럭이 같은 계근대를 번갈아 씁니다.
- 하역 자리, 선별 작업, 압착기, 엘리베이터, 반사로, 주조 라인은 “동시에 몇 대(몇 명)까지”가 정해져 있고, 자리가 없으면 앞에서 기다립니다.
- 파레트·제품 야적(버퍼) 이 가득 차면, 그 앞 단계가 멈추거나 줄어드는 식으로 막힐 수 있습니다. 실제 공장에서 창고가 꽉 찼을 때와 비슷합니다.
- 반사로 한 번의 용해(배치) 는 “정해진 파레트 개수가 모였을 때”만 시작합니다.
- 출하 쪽은 빈 트럭이 평균 간격을 두고 도착하되, 오전(08~12시)에 더 자주 오도록 잡혀 있습니다.  
  트럭은 재고를 최대 약 4시간까지 기다린 뒤, 채워진 만큼만 싣고 갑니다.
- 큐프레이크 vs SCR 같은 차종 선택은, 화면에서 정한 비율과 같은 확률로 고릅니다.
        """.strip()
    )

    st.markdown("##### 3) 이 화면(사이드바)에서만 바꿀 수 있는 것")
    st.success(
        "실행 버튼을 누를 때마다 아래만 사용자 입력으로 바뀌고, "
        "계근·하역·선별 시간, 압착 한 사이클 분, 용해·정련 병목(약 13시간) , 엘리베이터 속도, 야적 용량 등 나머지는 전부 기본 설계표입니다."
    )
    st.markdown(
        """
| 구간 | 이 화면에서 조절하는 것 |
|------|-------------------------|
| 기본 | 시뮬레이션 일수 |
| 입고 | 하루 트럭 대수, 트럭당 실은 무게, 하역 자리 수 |
| 선별·압착 | 선별 작업 인원(조), 압착기 대수, 파레트를 잠깐 쌓아 두는 공간(버퍼) 크기 |
| 용해·주조 | 반사로 대수, 한 배치당 톤수, 큐프레이크 vs SCR 비율 |
| 출하 | 빈 트럭이 오는 간격의 평균(분) — 실제로는 그 주변으로 들쭉날쭉 옵니다 |

배치 톤수와 파레트 수: 웹에서는 “파레트 하나 ≈ 2.5 t”로 고정해 두고, 배치 톤 ÷ 2.5로 파레트 개수를 맞춥니다.
        """.strip()
    )

    st.markdown("##### 4) 설정 파일에 있는데, 지금 시뮬에서는 안 쓰거나 다르게 쓰는 것")
    st.warning(
        "문서·설정에 이름이 있어도 실제 돌아가는 코드에서 참조하지 않는 값이 있습니다. "
        "결과가 ‘설명 문서와 숫자가 왜 다르지?’처럼 느껴질 때는 아래를 보면 됩니다."
    )
    st.markdown(
        """
- 산더미(한 덩이) 크기, 지게차 대수, 용해 가능 상한 톤수  
  → 현재 시뮬레이션 계산에는 거의 쓰이지 않습니다.
- 선별 시간  
  → 설정에 쪼개진 분 단위 값이 있어도, 실제 동작은 트럭당 30분으로 고정되어 있습니다.
- 주조에서 비율로 나눈 뒤 단위 개수  
  → 소수점을 버리면서 설정 비율과 아주 조금 어긋날 수 있습니다.
        """.strip()
    )

    st.markdown("##### 5) CP-SAT(터미널 보조) 한 줄 요약")
    st.markdown(
        """
각 배치는 한 대의 반사로만 쓰고, 같은 반사로에서는 시간이 겹치지 않게 놓되,
“준비된 시각 이후에만 시작” 같은 조건을 지키면서 전체가 끝나는 시각을 가능한 한 앞당기는 식의 계산입니다.
배치 한 번이 걸리는 시간은 엘리베이터·준비·용해·주조 중 긴 쪽 주조 시간으로 단순화해 잡습니다.  
그래서 이 웹에서 보는 시간 흐름(SimPy) 과 숫자가 완전히 같지는 않을 수 있습니다.
        """.strip()
    )
    if _SIM_INPUTS_CONSTRAINTS_MD.is_file():
        st.caption(f"검증·갱신용 원문: 저장소 `{_SIM_INPUTS_CONSTRAINTS_MD.relative_to(_REPO_ROOT)}`")

    with st.expander("원문 마크다운 전체 보기 (표·필드명)", expanded=False):
        md = _read_markdown_file(_SIM_INPUTS_CONSTRAINTS_MD)
        if md.strip():
            st.markdown(md)
        else:
            st.write("파일이 비어 있거나 읽을 수 없습니다.")


def _render_layperson_plotly_figures(
    metrics: Metrics,
    cfg: SimulationConfig,
    analysis: Analysis,
    *,
    chart_key_prefix: str = "gunsan_layperson",
) -> None:
    """일반인용 Plotly 6종(가동률·버퍼·트럭·체류·일별·간트)을 동일 레이아웃으로 렌더."""
    _lp_figs = build_layperson_visual_figures(metrics, cfg, analysis)
    _lp_r1a, _lp_r1b = st.columns(2)
    with _lp_r1a:
        st.plotly_chart(
            _style_layperson_result_figure(_lp_figs["utilization"]),
            use_container_width=True,
            key=f"{chart_key_prefix}_util",
        )
    with _lp_r1b:
        st.plotly_chart(
            _style_layperson_result_figure(_lp_figs["buffers"]),
            use_container_width=True,
            key=f"{chart_key_prefix}_buffer",
        )
    _lp_r2a, _lp_r2b = st.columns(2)
    with _lp_r2a:
        st.plotly_chart(
            _style_layperson_result_figure(_lp_figs["trucks"]),
            use_container_width=True,
            key=f"{chart_key_prefix}_trucks",
        )
    with _lp_r2b:
        st.plotly_chart(
            _style_layperson_result_figure(_lp_figs["lead_times"]),
            use_container_width=True,
            key=f"{chart_key_prefix}_lead",
        )
    _lp_r3a, _lp_r3b = st.columns(2)
    with _lp_r3a:
        st.plotly_chart(
            _style_layperson_result_figure(_lp_figs["daily_output"]),
            use_container_width=True,
            key=f"{chart_key_prefix}_daily",
        )
    with _lp_r3b:
        st.plotly_chart(
            _style_layperson_result_figure(_lp_figs["furnace_gantt"]),
            use_container_width=True,
            key=f"{chart_key_prefix}_gantt",
        )


def _render_glossary_page() -> None:
    """웹 대시보드 용어·약어 페이지(파일 기반 + 편집)."""
    st.subheader("📚 용어약어")
    st.caption(
        f"원본 파일: `{_TERMS_GLOSSARY_MD.relative_to(_REPO_ROOT)}`. "
        "저장하면 같은 서버를 보는 다른 사용자도 즉시 같은 내용을 확인할 수 있습니다."
    )
    if not _TERMS_GLOSSARY_MD.is_file():
        st.warning(
            f"`{_TERMS_GLOSSARY_MD.relative_to(_REPO_ROOT)}` 파일이 없어 기본 템플릿을 표시합니다. "
            "편집 후 저장하면 파일이 생성됩니다."
        )
    glossary_md = _read_markdown_file(
        _TERMS_GLOSSARY_MD,
        default_text=(
            "# 용어 및 약어\n\n"
            "아래 내용을 Markdown으로 자유롭게 편집하세요.\n\n"
            "## 약어\n"
            "- SCR: South Wire Rod\n"
            "- DES: Discrete Event Simulation\n"
            "- CP-SAT: Constraint Programming + SAT\n"
        ),
    )
    st.markdown(glossary_md)

    with st.expander("✍️ 용어/약어 편집", expanded=False):
        edited_glossary_md = st.text_area(
            "용어/약어 Markdown",
            value=glossary_md,
            height=520,
            key="terms_glossary_editor",
        )
        if st.button("💾 용어/약어 저장", type="primary", key="save_terms_glossary"):
            _save_markdown_file(_TERMS_GLOSSARY_MD, edited_glossary_md)
            st.success("용어/약어 내용을 저장했습니다. 다른 사용자도 새로고침 후 동일 내용을 볼 수 있습니다.")
            st.rerun()


def _render_doc_process_detail() -> None:
    """실행 전·문서 내비: 공정 상세."""
    st.markdown("## 군산 공장 하이브리드 공정 상세")
    st.markdown(
        "스크랩 구리 입고부터 완제품 출하까지 **5단계** 공정의 상세 설명입니다. "
        "각 단계는 **소제목 → 요약 불릿** 순으로 읽기 쉽게 정리되어 있습니다."
    )
    process_md = _read_markdown_file(_PROCESS_DETAIL_MD)
    if process_md:
        st.markdown(process_md)
    else:
        st.warning("공정 상세 파일이 비어 있습니다. 아래 편집기에서 내용을 작성해 저장해 주세요.")

    with st.expander("✍️ 공정 상세 편집", expanded=False):
        edited_process_md = st.text_area(
            "공정 상세 Markdown",
            value=process_md,
            height=620,
            key="process_detail_editor",
        )
        if st.button("💾 공정 상세 저장", type="primary", key="save_process_detail"):
            _save_markdown_file(_PROCESS_DETAIL_MD, edited_process_md)
            st.success("공정 상세를 저장했습니다. 다른 사용자도 새로고침 후 동일 내용을 확인할 수 있습니다.")
            st.rerun()


def _render_doc_methodology() -> None:
    """실행 전·문서 내비: 방법론 및 라이브러리."""
    st.markdown("## 시뮬레이션 방법론 및 사용 라이브러리")
    st.markdown("""
    본 시뮬레이션은 학술적으로 검증된 방법론과 산업 표준 라이브러리를 활용하여
    결과의 신뢰성과 재현성을 보장합니다.
    """)

    _render_simpy_cpsat_overview_for_dashboard()

    # SimPy 설명
    st.markdown("---")
    st.markdown("### 1. SimPy - 이산사건 시뮬레이션 (Discrete Event Simulation)")

    col_simpy1, col_simpy2 = st.columns([2, 1])
    with col_simpy1:
        st.markdown("""
        SimPy는 Python 기반 이산사건 시뮬레이션(DES) 프레임워크로,
        2002년 최초 출시 이후 20년 이상 학술 및 산업 분야에서 검증되었습니다.

        #### 학술적/산업적 신뢰성
        - Google Scholar: 수천 편의 학술 논문에서 인용
        - 적용 분야: 제조업 공정, 물류/공급망, 의료 시스템, 통신 네트워크
        - 글로벌 기업: Boeing, Toyota, DHL 등의 시뮬레이션 프로젝트에 활용

        #### 이산사건 시뮬레이션(DES)이란?
        연속 시간을 모사하지 않고 이벤트 발생 시점만 처리하여 계산 효율을 극대화하는 방법론입니다.
        제조업 공정 시뮬레이션의 국제 표준 방법론으로 인정받고 있습니다.
        """)
    with col_simpy2:
        st.markdown("""
        | 항목 | 내용 |
        |------|------|
        | 라이선스 | MIT (오픈소스) |
        | 버전 | 4.1+ |
        | 최초 출시 | 2002년 |
        | 유지보수 | 활발 (지속 업데이트) |
        """)

    with st.expander("💡 본 프로젝트에서 SimPy 활용 상세"):
        st.markdown("""
        ```python
        # 자원 경쟁 모델링 - 대기열 및 선착순 처리 자동 관리
        self.weighbridge = simpy.Resource(env, capacity=1)   # 계근대 1개
        self.furnaces = simpy.Resource(env, capacity=2)       # 반사로 2개

        # 버퍼 관리 - 용량 초과 시 생산 라인 자동 정지
        self.pallet_buffer = simpy.Store(env, capacity=160)   # 파레트 버퍼

        # 병렬 프로세스 - 큐프레이크/SCR 동시 주조
        yield self.env.all_of([flake_proc, scr_proc])
        ```

        모델링된 자원:
        - 계근대(1개), 하역장(2개), 압착기(1개), 반사로(2개), 엘리베이터(1개)
        - 파레트 버퍼(160개), 큐프레이크 야적(100포대), SCR 코일 야적(75코일)
        """)

    # CP-SAT 설명
    st.markdown("---")
    st.markdown("### 2. Google OR-Tools CP-SAT - 제약 만족 프로그래밍 최적화")

    col_cpsat1, col_cpsat2 = st.columns([2, 1])
    with col_cpsat1:
        st.markdown("""
        CP-SAT (Constraint Programming - SAT Solver)는 Google Research의
        Operations Research Team이 개발한 최적화 솔버입니다.

        #### 학술적/산업적 신뢰성
        - MiniZinc Challenge: 국제 제약 프로그래밍 경진대회에서 지속적 상위권 기록
        - Google 내부 활용: 자원 배분, 직원 스케줄링, 광고 최적화에 실전 적용
        - 학술 검증: 수천 편의 논문에서 벤치마크 솔버로 활용

        #### 핵심 특징
        - 최적성 증명: 최적해 발견 시 "더 나은 해가 없음"을 수학적으로 증명
        - 작업 스케줄링 특화: `IntervalVar`, `NoOverlap` 등 스케줄링 전용 기능 제공
        """)
    with col_cpsat2:
        st.markdown("""
        | 항목 | 내용 |
        |------|------|
        | 개발사 | Google Research |
        | 라이선스 | Apache 2.0 (오픈소스) |
        | 버전 | 9.10+ |
        | 솔버 유형 | SAT + CP 하이브리드 |
        """)

    with st.expander("💡 본 프로젝트에서 CP-SAT 활용 상세"):
        st.markdown("""
        반사로 배치 스케줄 최적화 문제:

        ```python
        # 변수 정의
        start = model.NewIntVar(release_min, horizon_min, f"start_{batch_id}")

        # 제약 조건: 같은 반사로 내 작업 중첩 금지
        model.AddNoOverlap(intervals_per_furnace[f])

        # 목적 함수: 메이크스팬 최소화
        model.Minimize(makespan)
        ```

        최적화 문제 구조:
        - 변수: 배치 시작 시각, 반사로 배정 (1 또는 2)
        - 제약: 파레트 32개 준비 후 시작, 동일 반사로 작업 비중첩
        - 목적: 전체 완료 시간(Makespan) 최소화

        결과 해석:
        - `OPTIMAL` 상태 시: 해당 메이크스팬이 이론적 최선임을 보장
        - `FEASIBLE` 상태 시: 실행 가능한 해이나 최적성 미증명
        """)

    # Matplotlib 설명
    st.markdown("---")
    st.markdown("### 3. Matplotlib - 시각화 및 애니메이션")

    st.markdown("""
    Matplotlib은 Python 시각화의 사실상 표준(de facto standard)으로,
    과학/공학 분야에서 가장 널리 사용되는 플로팅 라이브러리입니다.

    | 항목 | 내용 |
    |------|------|
    | 라이선스 | PSF License (Python Software Foundation) |
    | 버전 | 3.8+ |
    | 활용 | 공장 레이아웃 애니메이션, 버퍼 시계열 그래프, GIF/MP4 출력 |
    """)

    # Plotly / Streamlit 설명
    st.markdown("---")
    st.markdown("### 4. Plotly & Streamlit - 인터랙티브 대시보드")

    col_ui1, col_ui2 = st.columns(2)
    with col_ui1:
        st.markdown("""
        Plotly
        - 인터랙티브 차트 라이브러리
        - 줌, 팬, 호버 등 동적 기능 지원
        - 버퍼 시계열, Gantt 차트, 히스토그램 렌더링
        """)
    with col_ui2:
        st.markdown("""
        Streamlit
        - Python 기반 웹 앱 프레임워크
        - 데이터 과학/ML 대시보드에 최적화
        - 실시간 파라미터 조정 및 즉시 결과 확인
        """)

    # 방법론 요약
    st.markdown("---")
    st.markdown("### 📊 시뮬레이션 결과 신뢰성 요약")

    st.markdown("""
    | 구분 | 방법론 | 신뢰성 근거 |
    |------|--------|-------------|
    | 공정 시뮬레이션 | 이산사건 시뮬레이션 (DES) | 제조업 국제 표준, 20년+ 검증된 SimPy |
    | 스케줄 최적화 | 제약 만족 프로그래밍 (CP-SAT) | Google 개발, 국제 대회 검증, 최적성 수학적 증명 |
    | 불확실성 반영 | 지수분포 기반 확률 모델 | 도착 과정의 표준 확률 모델 (출하 트럭) |
    | 재현성 | 랜덤 시드 고정 | 동일 시드로 동일 결과 보장 |
    """)

    with st.expander("⚠️ 결과 해석 시 주의사항"):
        st.markdown("""
        1. 확정적 가정: 작업 시간(용해·정련 병목 약 13시간, 압착 90초 등)은 고정값으로 모델링
           - 실제 변동성 반영 필요 시 확률 분포 적용 가능

        2. 단순화된 설비 모델: 설비 고장, 유지보수 일정 미반영
           - 추후 확장 가능

        3. 시드 기반 재현성: `random_seed` 고정으로 재현성 확보
           - 다른 시드로 반복 실험하여 통계적 신뢰구간 산출 권장

        4. 입력 데이터 의존성: 파라미터 값의 정확도가 결과 품질에 직접 영향
           - 현장 데이터 기반 파라미터 검증 필요
        """)


def _render_doc_inputs_layperson() -> None:
    """실행 전·문서 내비: 입력과 규칙(쉬운 설명)."""
    st.markdown(
        "아래는 저장소 문서 `docs/simulation_inputs_constraints.md`와 같은 내용을 "
        "표·전문 용어를 줄여 풀어 쓴 전체 안내입니다. "
        "요약만 보려면 **실행·결과** 화면의 「이 시뮬 숫자가 어디서 오나요?」 접기 메뉴를 이용하세요."
    )
    _render_simulation_inputs_layperson_guide()


# 접근 제어(로그인 페이지)
_ACCESS_PASSWORD = "6501"
_DASHBOARD_AUTH_COOKIE = "gunsan_scr_dashboard"
_DASHBOARD_AUTH_TOKEN = hashlib.sha256(_ACCESS_PASSWORD.encode()).hexdigest()
_COOKIE_CM_KEY = "gunsan_dashboard_cookie_mgr"
_ENABLE_PASSWORD_AUTH = os.getenv("GUNSAN_ENABLE_PASSWORD_AUTH", "0") == "1"


def _auth_cookie_manager() -> CookieManager:
    """매 rerun마다 새 인스턴스를 만든다.

    `CookieManager` 는 내부에서 Streamlit 커스텀 컴포넌트를 호출하기 때문에
    `@st.cache_resource` 로 감싸면 ``CachedWidgetWarning`` 이 발생한다.
    컴포넌트는 ``key`` 기준으로 dedup 되므로 매 rerun 재생성해도 안전하다.
    """
    return CookieManager(key=_COOKIE_CM_KEY)


def _build_process_graphviz(
    trucks_per_day: int,
    payload_ton: float,
    flake_ratio: int,
    furnace_count: int,
    bottleneck: str | None = None,
) -> str:
    """가독성 중심 상세 공정 Graphviz(상단: 입고~엘리베이터, 하단: 용해·주조·출하)."""
    total_inbound_ton = trucks_per_day * payload_ton
    flake_ton = total_inbound_ton * (flake_ratio / 100.0)
    scr_ton = max(total_inbound_ton - flake_ton, 0.0)
    bottleneck_text = bottleneck or ""
    is_press = "압착" in bottleneck_text
    is_furnace = "반사로" in bottleneck_text
    is_outbound = "출하" in bottleneck_text

    sc = DEFAULT_CONFIG.sorting
    mc = DEFAULT_CONFIG.melting
    press_cycle = sc.forklift_load_min + sc.press_min_per_block + sc.pallet_stack_min
    el_cycle = mc.elevator_cycle_min
    melt_h = mc.melting_min / 60.0

    def style(is_bottle: bool, fill: str) -> str:
        border = "#dc2626" if is_bottle else "#5b6b7a"
        bg = "#fee2e2" if is_bottle else fill
        return f'shape=box style="rounded,filled" color="{border}" fillcolor="{bg}" penwidth=2'

    return f"""
digraph G {{
  rankdir=TB;
  graph [pad=0.45, nodesep=0.95, ranksep=1.45, bgcolor="white"];
  node [fontname="Malgun Gothic", fontsize=16, shape=box, style="rounded,filled", color="#5b6b7a", fillcolor="#eef6ff", penwidth=2, margin="0.32,0.24"];
  edge [color="#6b7280", penwidth=1.8, arrowsize=0.85];

  inbound   [label="트럭 입고\\n09~18시·오전80%\\n{trucks_per_day}대/일 · {total_inbound_ton:.0f}t/일"];
  weigh     [label="1차/2차 계근\\n각 5분"];
  unload    [label="하역\\n20분 · 베이 운영"];
  sorting   [label="선별\\n30분 · 8개 sub-pile"];
  press     [{style(is_press, "#eaf5ff")} label="압착/파레트\\n0.5t 사이클 {press_cycle:.1f}분\\n파레트 버퍼"];
  elevator  [label="엘리베이터\\n2파레트/{el_cycle:.0f}분"];
  furnace   [{style(is_furnace, "#fff4e8")} label="장입/용해\\n반사로 {furnace_count}대 · 병목 {melt_h:.1f}h"];
  casting   [label="하이브리드 주조\\nFlake {flake_ratio}% · SCR {100 - flake_ratio}%"];

  flake_yard [label="Flake 야적\\n{flake_ton:.0f}t", fillcolor="#e8f6ff"];
  scr_yard   [label="SCR 야적\\n{scr_ton:.0f}t", fillcolor="#ffeded"];
  flake_out  [{style(is_outbound, "#f4efff")} label="Flake 상차/출하\\n상차 → 계근 → 출고"];
  scr_out    [{style(is_outbound, "#f4efff")} label="SCR 상차/출하\\n상차 → 계근 → 출고"];

  {{ rank=same;
    inbound -> weigh -> unload -> sorting -> press -> elevator;
  }}
  elevator -> furnace;
  {{ rank=same;
    furnace -> casting;
  }}
  casting -> flake_yard;
  casting -> scr_yard;
  {{ rank=same;
    flake_yard -> flake_out;
    scr_yard -> scr_out;
  }}
}}
"""


def _process_flow_stage_is_bottleneck(stage_key: str, bottleneck: str | None) -> bool:
    """`analyze()` 가 넘기는 자원명(한글)과 공정 카드 키를 맞춘다."""
    if not bottleneck:
        return False
    bn = bottleneck
    mapping: dict[str, tuple[str, ...]] = {
        "truck": (),
        "weigh": ("계근",),
        "sort": ("선별",),
        "press": ("압착",),
        "elevator": ("엘리베이터",),
        "furnace": ("반사로",),
        "casting": (),
        "flake_out": ("출하",),
        "scr_out": ("출하",),
    }
    for needle in mapping.get(stage_key, ()):
        if needle in bn:
            return True
    return False


_PROCESS_FLOW_SECTION_CAPTION = (
    "슬라이더 값이 반영된 입고→출하 흐름입니다. "
    "시뮬 실행 후에는 가동률 기준 병목 자원에 해당하는 단계가 붉게 강조됩니다."
)


def _render_process_flow_section_header() -> None:
    """`st.subheader`/`st.caption` 대비 세로 여백을 줄인 섹션 머리글."""
    cap = html_module.escape(_PROCESS_FLOW_SECTION_CAPTION)
    st.markdown(
        f'<div class="gunsan-pf-section-wrap">'
        f"<h3>세부공정 프로세스</h3>"
        f'<p class="gunsan-pf-section-cap">{cap}</p>'
        f"</div>",
        unsafe_allow_html=True,
    )


def _render_process_flow_pipeline_html(
    trucks_per_day: int,
    payload_ton: float,
    flake_ratio: int,
    furnace_count: int,
    *,
    bottleneck: str | None = None,
    show_tooltips: bool = True,
    height: int | None = None,
) -> None:
    """세부 공정을 Plotly 대신 HTML 파이프라인 카드로 표시(다크 테마·슬라이더 값 반영)."""
    esc = html_module.escape
    total_inbound_ton = trucks_per_day * payload_ton
    flake_ton = total_inbound_ton * (flake_ratio / 100.0)
    scr_ton = max(total_inbound_ton - flake_ton, 0.0)

    sc = DEFAULT_CONFIG.sorting
    mc = DEFAULT_CONFIG.melting
    press_cycle = sc.forklift_load_min + sc.press_min_per_block + sc.pallet_stack_min
    el_cycle = mc.elevator_cycle_min
    melt_h = mc.melting_min / 60.0

    def card(
        key: str,
        title: str,
        desc: str,
        *,
        compact: bool = False,
    ) -> str:
        is_bn = _process_flow_stage_is_bottleneck(key, bottleneck)
        bn_cls = " gunsan-pf-card--bottleneck" if is_bn else ""
        sz = " gunsan-pf-card--compact" if compact else ""
        tip = f' title="{esc(desc)}"' if show_tooltips else ""
        badge = (
            '<span class="gunsan-pf-bn-badge" aria-label="병목 구간">병목</span>'
            if is_bn
            else ""
        )
        return (
            f'<article class="gunsan-pf-card{bn_cls}{sz}"{tip}>'
            f"{badge}"
            f'<h4 class="gunsan-pf-title">{esc(title)}</h4>'
            f'<p class="gunsan-pf-desc">{esc(desc)}</p>'
            f"</article>"
        )

    def conn() -> str:
        return '<span class="gunsan-pf-arrow" aria-hidden="true"></span>'

    main_cards: list[tuple[str, str, str]] = [
        ("truck", "트럭 입고", f"09~18h·오전 80% · {trucks_per_day}대/일 · {total_inbound_ton:.0f}t/일"),
        ("weigh", "계근/하역", "계근 5분 · 하역 20분"),
        ("sort", "선별", "30분 · 8 sub-pile"),
        ("press", "압착/파레트", f"0.5t {press_cycle:.1f}분/사이클 · 버퍼"),
        ("elevator", "엘리베이터", f"2파레트 / {el_cycle:.0f}분"),
        ("furnace", "장입/용해", f"반사로 {furnace_count}대 · 병목 약 {melt_h:.1f}h"),
        ("casting", "하이브리드 주조", f"Flake {flake_ratio}% · SCR {100 - flake_ratio}%"),
    ]
    main_html = conn().join(card(k, t, d) for k, t, d in main_cards)

    flake_track = "".join(
        [
            card("flake_yard", "Flake 야적", f"{flake_ton:.0f} t", compact=True),
            conn(),
            card("flake_out", "Flake 출하", "상차 → 계근 → 출고", compact=True),
        ]
    )
    scr_track = "".join(
        [
            card("scr_yard", "SCR 야적", f"{scr_ton:.0f} t", compact=True),
            conn(),
            card("scr_out", "SCR 출하", "상차 → 계근 → 출고", compact=True),
        ]
    )

    foot = (
        "상세 흐름: 입고 → 계근/하역 → 선별 → 압착/버퍼 → 엘리베이터 → "
        "용해(12h) → 주조 → 야적 → 상차/출하"
    )
    bn_note = ""
    if bottleneck:
        bn_note = (
            f'<p class="gunsan-pf-bn">병목 자원: <strong>{esc(bottleneck)}</strong> '
            f"(해당 단계는 <strong>병목</strong> 표기·붉은 강조로 구분됩니다)</p>"
        )

    html = f"""
<div class="gunsan-pf-root">
  <style>
    .gunsan-pf-root {{
      font-family: "Malgun Gothic", "Apple SD Gothic Neo", system-ui, sans-serif;
      background: rgba(15, 23, 42, 0.97);
      border: 1px solid rgba(51, 65, 85, 0.75);
      border-radius: 12px;
      padding: 0.5rem 0.65rem 0.4rem;
      color: #e2e8f0;
      box-sizing: border-box;
    }}
    .gunsan-pf-root * {{ box-sizing: border-box; }}
    .gunsan-pf-row {{
      display: flex;
      flex-wrap: wrap;
      align-items: stretch;
      justify-content: flex-start;
      gap: 0.15rem 0;
      row-gap: 0.5rem;
    }}
    .gunsan-pf-main {{
      flex-wrap: nowrap;
      align-items: stretch;
      justify-content: flex-start;
      overflow-x: auto;
      overflow-y: hidden;
      padding-bottom: 0.28rem;
      margin: 0 -0.15rem;
      padding-left: 0.15rem;
      padding-right: 0.15rem;
      border-bottom: 1px solid rgba(71, 85, 105, 0.45);
      margin-bottom: 0.28rem;
      scrollbar-width: thin;
      scrollbar-color: rgba(100,116,139,0.6) transparent;
    }}
    .gunsan-pf-main::-webkit-scrollbar {{
      height: 6px;
    }}
    .gunsan-pf-main::-webkit-scrollbar-thumb {{
      background: rgba(100, 116, 139, 0.55);
      border-radius: 4px;
    }}
    .gunsan-pf-main .gunsan-pf-card {{
      flex: 0 0 auto;
      width: clamp(5.65rem, 11.5vw, 8.85rem);
      max-width: none;
      min-width: 5.65rem;
    }}
    .gunsan-pf-main .gunsan-pf-arrow {{
      flex: 0 0 1rem;
    }}
    .gunsan-pf-card {{
      flex: 1 1 6.5rem;
      min-width: 5.75rem;
      max-width: 11rem;
      margin: 0.1rem 0.05rem;
      padding: 0.45rem 0.4rem 0.42rem;
      border-radius: 12px;
      background: linear-gradient(180deg, rgba(30,41,59,0.95) 0%, rgba(15,23,42,0.88) 100%);
      border: 1px solid color-mix(in srgb, var(--stage-accent) 55%, transparent);
      box-shadow: 0 4px 18px rgba(0,0,0,0.35), inset 0 1px 0 rgba(255,255,255,0.04);
      position: relative;
      overflow: hidden;
    }}
    .gunsan-pf-card::before {{
      content: "";
      position: absolute;
      left: 0; right: 0; top: 0;
      height: 3px;
      background: linear-gradient(90deg, var(--stage-accent), transparent);
      opacity: 0.9;
    }}
    .gunsan-pf-card--compact {{
      flex: 1 1 8rem;
      max-width: 14rem;
      padding: 0.38rem 0.35rem 0.36rem;
    }}
    .gunsan-pf-card--bottleneck {{
      border: 2px solid #ef4444 !important;
      background: rgba(69, 10, 10, 0.55) !important;
      box-shadow:
        0 0 0 1px rgba(252, 165, 165, 0.55),
        0 0 22px rgba(239, 68, 68, 0.35),
        0 4px 16px rgba(0,0,0,0.45) !important;
    }}
    .gunsan-pf-card--bottleneck::before {{
      height: 4px !important;
      background: linear-gradient(90deg, #ef4444, #f97316) !important;
      opacity: 1 !important;
    }}
    .gunsan-pf-card--bottleneck .gunsan-pf-title {{
      color: #fecaca;
      padding-right: 2.35rem;
    }}
    .gunsan-pf-card--bottleneck .gunsan-pf-desc {{
      color: #fca5a5;
    }}
    .gunsan-pf-bn-badge {{
      position: absolute;
      top: 0.32rem;
      right: 0.32rem;
      font-size: 0.56rem;
      font-weight: 800;
      letter-spacing: 0.04em;
      color: #450a0a;
      background: #fecaca;
      padding: 0.1rem 0.32rem;
      border-radius: 4px;
      line-height: 1.2;
      z-index: 1;
    }}
    .gunsan-pf-title {{
      margin: 0 0 0.25rem 0;
      font-size: 0.82rem;
      font-weight: 700;
      letter-spacing: -0.02em;
      color: #f1f5f9;
      line-height: 1.25;
    }}
    .gunsan-pf-card--compact .gunsan-pf-title {{ font-size: 0.78rem; }}
    .gunsan-pf-desc {{
      margin: 0;
      font-size: 0.68rem;
      line-height: 1.45;
      color: #94a3b8;
    }}
    .gunsan-pf-card--compact .gunsan-pf-desc {{ font-size: 0.64rem; }}
    .gunsan-pf-arrow {{
      flex: 0 0 1.1rem;
      align-self: center;
      width: 1.1rem;
      height: 1.1rem;
      margin: 0 0.05rem;
      flex-shrink: 0;
      opacity: 0.85;
      background: no-repeat center / contain
        url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='%2394a3b8' stroke-width='2' stroke-linecap='round' stroke-linejoin='round'%3E%3Cpath d='M5 12h14M13 6l6 6-6 6'/%3E%3C/svg%3E");
    }}
    .gunsan-pf-bridge {{
      display: flex;
      flex-direction: column;
      align-items: center;
      justify-content: center;
      gap: 0;
      padding: 0 0 0.15rem;
      margin-top: -0.2rem;
    }}
    .gunsan-pf-bridge-arrow {{
      width: 1.1rem;
      height: 1.15rem;
      opacity: 0.88;
      flex-shrink: 0;
      background: no-repeat center / contain
        url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='%2394a3b8' stroke-width='2' stroke-linecap='round' stroke-linejoin='round'%3E%3Cpath d='M12 5v14M6 13l6 6 6-6'/%3E%3C/svg%3E");
    }}
    .gunsan-pf-split-label {{
      text-align: center;
      font-size: 0.7rem;
      color: #64748b;
      letter-spacing: 0.12em;
      text-transform: uppercase;
      margin: 0 0 0.22rem;
    }}
    .gunsan-pf-split {{
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 0.42rem 0.75rem;
    }}
    @media (max-width: 720px) {{
      .gunsan-pf-split {{ grid-template-columns: 1fr; }}
    }}
    .gunsan-pf-track {{
      display: flex;
      flex-wrap: nowrap;
      align-items: stretch;
      justify-content: flex-start;
      gap: 0;
      overflow-x: auto;
      overflow-y: hidden;
      padding: 0.18rem 0.18rem;
      border-radius: 10px;
      background: rgba(15, 23, 42, 0.55);
      border: 1px dashed rgba(71, 85, 105, 0.55);
      scrollbar-width: thin;
      scrollbar-color: rgba(100,116,139,0.5) transparent;
    }}
    .gunsan-pf-track::-webkit-scrollbar {{
      height: 5px;
    }}
    .gunsan-pf-track::-webkit-scrollbar-thumb {{
      background: rgba(100, 116, 139, 0.45);
      border-radius: 4px;
    }}
    .gunsan-pf-track .gunsan-pf-card--compact {{
      flex: 0 0 auto;
      width: clamp(6.5rem, 28vw, 10.5rem);
      max-width: none;
    }}
    .gunsan-pf-track .gunsan-pf-arrow {{
      flex: 0 0 0.95rem;
    }}
    .gunsan-pf-foot {{
      margin: 0.18rem 0 0;
      font-size: 0.72rem;
      line-height: 1.5;
      color: #64748b;
      text-align: center;
    }}
    .gunsan-pf-bn {{
      margin: 0.2rem 0 0;
      font-size: 0.78rem;
      text-align: center;
      color: #fca5a5;
    }}
    .gunsan-pf-bn strong {{ color: #fecaca; }}
  </style>
  <div class="gunsan-pf-row gunsan-pf-main" role="list">
    {main_html}
  </div>
  <div class="gunsan-pf-bridge" aria-hidden="true">
    <span class="gunsan-pf-bridge-arrow"></span>
  </div>
  <div class="gunsan-pf-split-label">주조 이후 · 제품별 출하</div>
  <div class="gunsan-pf-split">
    <div class="gunsan-pf-track" role="group" aria-label="Flake 경로">{flake_track}</div>
    <div class="gunsan-pf-track" role="group" aria-label="SCR 경로">{scr_track}</div>
  </div>
  {bn_note}
  <p class="gunsan-pf-foot">{esc(foot)}</p>
</div>
"""
    # 고정 높이가 실제 레이아웃보다 크면 iframe 아래에 큰 빈 영역이 생김(탭 상단 등).
    iframe_h = height if height is not None else (452 if bottleneck else 400)
    components.html(html, height=iframe_h, scrolling=False)


def _build_process_timeline_figure(
    batch_ton: float,
    flake_ratio: int,
) -> go.Figure:
    """멋진 대시보드형 타임라인: 누적 네온 라인 + 도넛 + 히트맵."""

    def scenario_steps(ratio: int) -> tuple[list[str], list[float], float, float]:
        flake_ton = batch_ton * (ratio / 100.0)
        scr_ton = max(batch_ton - flake_ton, 0.0)
        press_for_batch = (batch_ton / 0.5) * 8.5
        elevator = (batch_ton / 80.0) * 160.0
        flake_cast = (flake_ton / 1.0) * 2.5 if flake_ton > 0 else 0.0
        scr_cast = (scr_ton / 4.0) * 10.0 if scr_ton > 0 else 0.0
        branch_elapsed = max(flake_cast, scr_cast)
        labels = ["1차 계근", "하역", "2차 계근", "선별", "압착/파레트", "엘리베이터", "장입 준비", "용해(12h)", "주조 셋업", "분기 주조 완료"]
        durations = [5, 20, 5, 30, press_for_batch, elevator, 120, 720, 90, branch_elapsed]
        return labels, durations, flake_ton, scr_ton

    low_ratio = max(10, flake_ratio - 20)
    high_ratio = min(90, flake_ratio + 20)
    scenarios = [("저 Flake", low_ratio), ("기준", flake_ratio), ("고 Flake", high_ratio)]

    base_labels, base_durations_min, base_flake_ton, base_scr_ton = scenario_steps(flake_ratio)
    base_durations_h = [v / 60.0 for v in base_durations_min]
    base_total_h = sum(base_durations_h)
    cumulative = []
    cur = 0.0
    for d in base_durations_h:
        cur += d
        cumulative.append(cur)

    fig = make_subplots(
        rows=2,
        cols=2,
        specs=[[{"type": "xy", "colspan": 2}, None], [{"type": "domain"}, {"type": "xy"}]],
        vertical_spacing=0.25,
        horizontal_spacing=0.12,
        subplot_titles=("기준 시나리오: 공정 누적시간 네온 타임라인", "공정 시간 비중", "경우의 수별 누적시간 히트맵"),
    )

    fig.add_trace(
        go.Bar(
            x=base_labels,
            y=base_durations_h,
            marker_color="rgba(56,189,248,0.35)",
            text=[f"{v:.2f}h" for v in base_durations_h],
            textposition="inside",
            hovertemplate="%{x}<br>소요: %{y:.2f}시간<extra></extra>",
            showlegend=False,
        ),
        row=1,
        col=1,
    )
    fig.add_trace(
        go.Scatter(
            x=base_labels,
            y=cumulative,
            mode="lines+markers+text",
            line=dict(color="#2563eb", width=4),
            marker=dict(size=10, color="#0ea5e9", line=dict(color="white", width=1)),
            text=[f"{v:.1f}h" for v in cumulative],
            textposition="top center",
            hovertemplate="%{x}<br>누적: %{y:.2f}시간<extra></extra>",
            showlegend=False,
        ),
        row=1,
        col=1,
    )

    fig.add_trace(
        go.Pie(
            labels=base_labels,
            values=base_durations_h,
            hole=0.58,
            sort=False,
            textinfo="percent",
            hovertemplate="%{label}<br>%{value:.2f}시간 (%{percent})<extra></extra>",
            marker=dict(line=dict(color="white", width=1)),
            showlegend=False,
        ),
        row=2,
        col=1,
    )

    heat_y = []
    heat_z = []
    for s_name, ratio in scenarios:
        labels, durations_min, _, _ = scenario_steps(ratio)
        cum_vals = []
        c = 0.0
        for d in durations_min:
            c += d / 60.0
            cum_vals.append(c)
        heat_y.append(f"{s_name} ({ratio}%)")
        heat_z.append(cum_vals)

    fig.add_trace(
        go.Heatmap(
            x=base_labels,
            y=heat_y,
            z=heat_z,
            colorscale="YlGnBu",
            colorbar=dict(title="누적시간(h)", len=0.8),
            hovertemplate="시나리오: %{y}<br>단계: %{x}<br>누적: %{z:.2f}시간<extra></extra>",
        ),
        row=2,
        col=2,
    )

    fig.update_layout(
        height=760,
        margin=dict(l=40, r=25, t=80, b=45),
        paper_bgcolor="white",
        plot_bgcolor="white",
        hoverlabel=_PLOTLY_HOVERLABEL,
        xaxis=dict(tickangle=-25),
        yaxis=dict(title="시간 (시간)", gridcolor="rgba(148,163,184,0.25)"),
        xaxis3=dict(tickangle=-25),
        yaxis3=dict(title="시나리오"),
    )
    fig.add_annotation(
        x=0.24,
        y=0.19,
        xref="paper",
        yref="paper",
        showarrow=False,
        text=f"<b>총 누적시간</b><br>{base_total_h:.2f}h",
        font=dict(size=14, color="#1e3a8a"),
    )
    fig.add_annotation(
        x=0.5,
        y=-0.08,
        xref="paper",
        yref="paper",
        showarrow=False,
        text=(
            f"기준 배치 {batch_ton:.0f}t | 기준 분기량 Flake {base_flake_ton:.1f}t, SCR {base_scr_ton:.1f}t "
            f"(분기 주조는 병렬 진행 기준으로 긴 쪽 시간 적용)"
        ),
        font=dict(size=12, color="#334155"),
    )
    return fig

# ---------------------------------------------------------------------------
# 페이지 설정
# ---------------------------------------------------------------------------

st.set_page_config(
    page_title="SCR공정 물류 시뮬레이션",
    page_icon="🏭",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# 접근 제어
# ---------------------------------------------------------------------------
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if _ENABLE_PASSWORD_AUTH:
    _cookie_mgr = _auth_cookie_manager()
    _cookie_val = _cookie_mgr.get(_DASHBOARD_AUTH_COOKIE)
    if (
        not st.session_state.authenticated
        and _cookie_val == _DASHBOARD_AUTH_TOKEN
    ):
        st.session_state.authenticated = True

    if (
        not st.session_state.authenticated
        and not st.session_state.get("_auth_cookie_bridge_done")
    ):
        st.session_state._auth_cookie_bridge_done = True
        st.rerun()

    if not st.session_state.authenticated:
        st.title("🔒 접근 제한")
        st.write("비밀번호를 입력해야 대시보드에 접속할 수 있습니다.")
        password_input = st.text_input("비밀번호", type="password", placeholder="비밀번호 입력")
        if st.button("접속", type="primary"):
            if password_input == _ACCESS_PASSWORD:
                st.session_state.authenticated = True
                _cookie_mgr.set(
                    _DASHBOARD_AUTH_COOKIE,
                    _DASHBOARD_AUTH_TOKEN,
                    key="gunsan_auth_cookie_set",
                    max_age=365.25 * 24 * 3600,
                    path="/",
                    same_site="lax",
                )
                st.rerun()
            else:
                st.error("비밀번호가 올바르지 않습니다.")
        st.stop()
else:
    st.session_state.authenticated = True

# ---------------------------------------------------------------------------
# 사이드바 - 파라미터 입력
# ---------------------------------------------------------------------------

if st.sidebar.button(
    "🏠 홈",
    key="gunsan_sidebar_home",
    use_container_width=True,
    help="브라우저 새로고침(F5)과 같습니다. 세션·위젯 상태가 초기화되고 페이지가 다시 로드됩니다.",
):
    # F5와 동일: 전체 탭 새로고침(세션 상태·사이드바 기본값까지 초기화)
    components.html(
        "<script>window.top.location.reload();</script>",
        height=0,
        width=0,
    )

st.sidebar.title("🏭 시뮬레이션 파라미터")
st.sidebar.caption(
    "ℹ️ 각 항목 옆 ⓘ 버튼과 결과 차트의 마우스 호버로 상세 정보를 볼 수 있습니다."
)
run_button = st.sidebar.button(
    "🚀 시뮬레이션 실행", type="primary", use_container_width=True,
    help="""▶️ 시뮬레이션 실행

설정한 파라미터로 이산사건 시뮬레이션을 실행합니다.

실행 내용:
1. SimPy 기반 DES 시뮬레이션 실행
2. 공정별 이벤트 로그 생성
3. KPI 및 병목 분석
4. 시각화 결과 생성

예상 소요 시간:
- 7일 시뮬레이션: 약 1~2초
- 30일 시뮬레이션: 약 3~5초

결과 확인:
- 핵심 지표 (KPI 카드)
- 병목 진단 및 권장사항
- 자원 가동률 차트
- 버퍼 시계열 그래프
- 반사로 Gantt 차트
- 트럭 흐름 분석"""
)
st.sidebar.caption("필요할 때 언제든 위 버튼으로 즉시 실행하세요.")

st.sidebar.header("1. 기본 설정")
st.sidebar.caption("시뮬레이션 실행 기간 설정")
sim_days = st.sidebar.slider(
    "시뮬레이션 일수", 1, 30, DEFAULT_CONFIG.sim_days,
    help="""📅 시뮬레이션 기간 설정

시뮬레이션을 실행할 총 일수를 지정합니다.

권장 설정:
- 단기 분석: 1~3일 (빠른 테스트)
- 중기 분석: 7~14일 (정상 상태 도달)
- 장기 분석: 14~30일 (통계적 신뢰성 확보)

주의사항:
- 일수가 길수록 시뮬레이션 실행 시간 증가
- 반사로 용해·정련 병목(약 13시간) 주기 고려 시 최소 3일 권장
- 초기 warm-up 기간(1~2일) 후 정상 상태 도달"""
)
random_seed = DEFAULT_CONFIG.random_seed

st.sidebar.header("2. 입고/하역")
st.sidebar.caption("스크랩 원료 트럭 입고 및 하역 설정")
trucks_per_day = st.sidebar.slider(
    "일일 트럭 대수", 1, 20, DEFAULT_CONFIG.inbound.trucks_per_day,
    help="""🚛 일일 스크랩 입고 트럭 대수

하루 동안 공장에 도착하는 스크랩 원료 운반 트럭의 평균 대수입니다.

입고 프로세스:
1. 트럭 도착 → 1차 계근 (총중량 측정)
2. 하역장 이동 → 스크랩 하역 (20~30분)
3. 2차 계근 (공차 중량) → 출차

일일 입고량 계산:
- 총 입고량 = 트럭 대수 × 트럭당 적재량
- 예: 10대 × 20t = 200t/일

파라미터 영향:
- 증가 시: 일일 원료 공급량 증가, 하역장/선별 부하 증가
- 감소 시: 원료 부족으로 반사로 가동률 저하 가능

실제 공장 고려사항(소재 공장 운영 정보 기준):
- 입고 창구 09~18시, 그중 약 80%는 오전(12시 이전)에 균등 분포로 도착
- 피크 시간대 하역장·계근대 혼잡 가능성"""
)
payload_ton = st.sidebar.slider(
    "트럭당 적재량 (t)", 10.0, 30.0, DEFAULT_CONFIG.inbound.payload_ton, 1.0,
    help="""⚖️ 트럭 1대당 스크랩 적재량 (톤)

각 트럭이 운반하는 스크랩 원료의 평균 중량입니다.

적재량 구성:
- 총중량(1차 계근) - 공차중량(2차 계근) = 순 적재량
- 일반적인 범위: 15~25톤/대

후속 공정 영향:
- 트럭 1대 하역 → 1개 더미 생성
- 1개 더미 → 8개 sub-pile로 선별
- sub-pile → 파레트로 압착

일일 처리량 계산:
- 일일 입고량 = 트럭 대수 × 적재량
- 예: 10대 × 20t = 200t/일

병목 관계:
- 반사로 1배치 = 30~80톤 (설정에 따라)
- 입고량 > 처리량 시 야적장 재고 증가
- 입고량 < 처리량 시 반사로 유휴 발생"""
)
unloading_bays = st.sidebar.slider(
    "하역 베이 수", 1, 4, DEFAULT_CONFIG.inbound.unloading_bays,
    help="""🏗️ 하역장 동시 작업 가능 대수

하역장에서 동시에 하역 작업을 수행할 수 있는 트럭 수입니다.

하역 작업 상세:
- 소요 시간: 약 20~30분/대
- 작업 내용: 덤프 트럭으로 스크랩 투하
- 결과물: 선별 대기 더미 1개/트럭

용량 영향:
- 1 베이: 순차 처리, 트럭 대기 시간 발생
- 2 베이: 병렬 처리, 처리량 2배
- 3+ 베이: 트럭 도착률 고려 필요

병목 진단:
- 하역 베이 가동률 > 90%: 병목 가능성
- 트럭 평균 대기시간 증가 시 베이 추가 검토

설비 비용 고려:
- 베이 추가는 초기 투자비 증가
- ROI 계산: 트럭 대기비용 vs 설비 투자비"""
)

st.sidebar.header("3. 선별/압착")
st.sidebar.caption("스크랩 등급 선별 및 파레트 압착 설정")
sorters = st.sidebar.slider(
    "선별 작업조 수", 1, 4, DEFAULT_CONFIG.sorting.sorters,
    help="""👷 선별 작업조(팀) 수

스크랩 원료를 등급별로 분류하고 이물질을 제거하는 작업조 수입니다.

선별 작업 내용:
- 스크랩 품질/등급 분류
- 비금속 이물질 제거 (플라스틱, 고무 등)
- 크기별 분류
- 소요 시간: 약 30분/더미

작업 흐름:
- 트럭 1대 하역 → 1개 더미
- 1개 더미 → 선별 후 8개 sub-pile
- sub-pile → 압착기로 이동

용량 영향:
- 작업조 증가 → 선별 처리량 증가
- 병렬 작업 가능 (각 조가 독립적으로 작업)

인력 비용 고려:
- 작업조당 2~3명 인력 필요
- 교대 근무 시 조 수 × 교대 수 인력 필요
- 선별 품질과 속도의 트레이드오프"""
)
press_machines = st.sidebar.slider(
    "압착기 대수", 1, 4, DEFAULT_CONFIG.sorting.press_machines,
    help="""🔧 압착기 설비 대수

선별된 sub-pile을 파레트로 압착하는 압착기의 대수입니다.

압착 공정 상세:
- 입력: 선별된 sub-pile (약 0.5톤)
- 출력: 압착 파레트 (규격화된 블록)
- 사이클 타임: 지게차+압착(90초)+파레트 적재(덩어리당 약 3분)를 블록당 반복

처리량 계산:
- 1대 기준: 약 6 파레트/시간 전후(설정값에 따라 변동)
- 일 처리량: 7 × 24 = 168 파레트/일 (연속 가동 시)

병목 영향:
- 압착기가 병목인 경우:
  - sub-pile 대기 큐 증가
  - 파레트 생성 지연
  - 반사로 장입 대기 발생

투자 판단 기준:
- 압착기 가동률 > 85%: 추가 투자 검토
- sub-pile 대기 큐 지속 증가: 병목 신호
- 파레트 버퍼 자주 비어있음: 압착기 부족"""
)
pallet_buffer_capacity = st.sidebar.slider(
    "파레트 버퍼 용량", 50, 300, DEFAULT_CONFIG.sorting.pallet_buffer_capacity, 10,
    help="""📦 파레트 버퍼 최대 적재량

압착 완료된 파레트가 반사로 장입 전까지 대기하는 버퍼의 용량입니다.

버퍼 역할:
- 압착 공정과 용해 공정 간 디커플링
- 공정 간 속도 차이 흡수
- 반사로 장입 대기열 역할

용량 산정 기준:
- 반사로 1배치(80 t) = 파레트 32개 전후(배치 톤수 ÷ 2.5 t)
- 최소 권장: 1배치 분량 이상
- 권장: 2일치 야적(문서 기준 최대 160개 전후)과 균형

용량 부족 시:
- 버퍼 full → 압착기 정지
- 압착기 정지 → sub-pile 대기 증가
- 연쇄적 공정 지연 발생

용량 과다 시:
- 불필요한 공간 점유
- 재고 관리 비용 증가
- 파레트 품질 저하 가능 (장기 보관 시)

모니터링 포인트:
- 평균 점유율: 50~70% 적정
- 최대 점유율: 90% 이하 유지 권장"""
)

st.sidebar.header("4. 용해/주조")
st.sidebar.caption("반사로 용해 및 제품 주조 설정 (핵심 병목)")
furnace_count = st.sidebar.slider(
    "반사로 대수", 1, 3, DEFAULT_CONFIG.melting.furnace_count,
    help="""🔥 반사로(Reverberatory Furnace) 대수

스크랩을 고온으로 용해하여 용탕(molten copper)을 만드는 반사로의 대수입니다.

반사로 운영 사이클(문서 기준 병목):
1. 장입·사전 준비: 엘리베이터 왕복·투입 준비(설정: 약 2시간 전후)
2. 용해·정련 병목: 8시간 용해 + 산화·슬래깅 각 30분 + 환원 4시간 → 약 13시간 동안 신규 장입 불가
3. 주조 (Casting): 홀딩로 셋업 후 큐플레이크·SCR 병렬 생산
4. 준비: 다른 반사로는 병렬로 다음 배치 준비 가능(2기 운영 전제)

총 사이클 타임: 약 24시간 운전 관점과 맞물리도록 조정됨

처리량 계산:
- 1대: ~30톤/일 (24시간 사이클)
- 2대: ~60톤/일 (교대 운용)
- 3대: ~90톤/일 (연속 생산)

병목 분석:
- 용해·산화·슬래깅·환원 구간이 핵심 병목(시뮬에서는 통합 분으로 모델링)
- 반사로 추가가 처리량 증대의 가장 직접적 방법
- 단, 대규모 설비 투자 필요

에너지 비용:
- 가스 버너 연료비 (주요 운영 비용)
- 용해 온도: 1,100~1,200°C
- 전력비: 버너, 제어 시스템 등"""
)
batch_ton = st.sidebar.slider(
    "배치 단위 (t)", 40.0, 200.0, DEFAULT_CONFIG.melting.batch_ton, 10.0,
    help="""⚗️ 반사로 1회 배치 용량 (톤)

반사로에 한 번에 장입하여 용해하는 스크랩의 총 중량입니다.

배치 구성:
- 배치 톤수 ÷ 파레트당 중량 = 필요 파레트 수
- 예: 30톤 ÷ 0.5톤 = 60 파레트/배치

배치 크기 영향:

작은 배치 (40~80톤):
- 장점: 빠른 사이클, 유연한 운영
- 단점: 배치당 고정 시간 비중 증가

큰 배치 (100~200톤):
- 장점: 배치당 효율 증가
- 단점: 긴 사이클, 파레트 대기 시간 증가

최적화 고려사항:
- 파레트 버퍼 용량과의 균형
- 일일 입고량 대비 배치 횟수
- 주조 라인 용량과의 매칭

실제 운영 팁:
- 일 1~2배치가 일반적
- 야간 용해 → 주간 주조 패턴 고려"""
)
flake_ratio = st.sidebar.slider(
    "큐프레이크 비율 (%)", 0, 100, int(DEFAULT_CONFIG.casting.flake_ratio * 100),
    help="""🥧 주조 제품 비율: 큐프레이크 vs SCR

용탕을 주조할 때 큐프레이크와 SCR 코일의 생산 비율입니다.

제품 유형:

큐프레이크 (Cu Flake):
- 형태: 얇은 구리 플레이크/칩
- 단위: 25kg 포대
- 용도: 전해 정련, 합금 제조
- 주조 속도: 약 1톤/3.1분 (포대 1개)

SCR 코일 (South Wire Rod):
- 형태: 연속 주조 구리봉 (코일)
- 단위: 4톤/코일
- 용도: 전선, 케이블 제조
- 주조 속도: 약 4톤/12.5분 (코일 1개)

비율 설정 가이드:
- 운영 정보 기본값: 큐플레이크 20% / SCR 80%
- 50/50: 두 품목 수요가 비슷할 때
- 큐플레이크 비중을 높일 때: 예) 40/60
- SCR 비중을 높일 때: 예) 10/90

주의사항:
- 두 라인 병렬 운영 (동시 주조)
- 야적장 용량 고려 필요
  - 큐프레이크: 100포대 버퍼
  - SCR: 75코일 버퍼
- 출하 트럭 수요와 매칭 필요"""
)

st.sidebar.header("5. 출하")
st.sidebar.caption("완제품 야적 및 출하 트럭 설정")
empty_truck_interval = st.sidebar.slider(
    "출하 트럭 평균 간격 (분)", 30, 180, int(DEFAULT_CONFIG.outbound.empty_truck_interval_min),
    help="""🚚 출하 빈 트럭 도착 평균 간격 (분)

완제품(큐프레이크/SCR)을 실어갈 빈 트럭이 도착하는 평균 시간 간격입니다.

출하 프로세스:
1. 빈 트럭 도착 → 대기열 진입
2. 야적장에서 제품 상차 (30~60분)
3. 2차 계근 (적재 중량 확인)
4. 출고

도착 간격 모델:
- 평균 간격은 지수분포로 전진한 뒤, 08~12시(80%) 또는 12~18시(20%) 창에서 시각을 잡는 방식으로 오전 출하 비중을 반영

간격 설정 가이드:

짧은 간격 (30~60분):
- 활발한 출하 수요
- 야적장 재고 빠르게 소진
- 트럭 대기 시간 발생 가능

긴 간격 (120~180분):
- 완만한 출하 수요
- 야적장 재고 누적 가능
- 야적장 용량 초과 주의

균형 포인트:
- 생산량 ≈ 출하량 유지
- 야적장 점유율 50~70% 목표
- 트럭 평균 대기시간 최소화"""
)

_inject_dashboard_base_styles()
_inject_streamlit_help_tooltip_styles(show_widget_help=True)

# ---------------------------------------------------------------------------
# 메인 — 제목·안내·탭(시뮬레이션 / 용어약어)
# ---------------------------------------------------------------------------

with st.container(border=False):
    st.title("🏭 SCR공정 물류 시뮬레이션")
    st.markdown(
        f"""
        <div style="text-align:right; color:#94a3b8; font-size:12px; margin-top:-6px;">
        {_BUILD_INFO_TEXT}
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("""
스크랩 구리 입고(09~18h·오전 80%) → 선별/압착 → 장입/용해(병목 약 13h) → 하이브리드 주조 → 완제품 출하(20t·오전 80%)의
5단계 공정을 SimPy 이산사건 시뮬레이션으로 모델링합니다.
왼쪽 사이드바에서 파라미터를 조정한 뒤 시뮬레이션 실행 버튼을 누르세요.
""")
    st.info(
        "자유롭게 여러 번 실행해 보셔도 됩니다. "
        "이 화면의 시뮬레이션은 실제 공장 설비나 외부 운영 시스템과 연결되어 있지 않으며, "
        "슬라이더를 바꾸고 실행 버튼을 반복해서 눌러도 생산·물류 현장에 직접적인 영향을 주지 않습니다. "
        "계산은 이 앱이 돌아가는 PC(또는 내부에 배포한 Streamlit 서버) 안에서만 이루어지며, "
        "보통 수 초 안에 끝나는 가벼운 수준이라 설정을 바꿔 가며 시험하기에 부담이 적습니다."
    )
    if run_button:
        inbound = dataclasses.replace(
            DEFAULT_CONFIG.inbound,
            trucks_per_day=trucks_per_day,
            payload_ton=payload_ton,
            unloading_bays=unloading_bays,
        )
        sorting = dataclasses.replace(
            DEFAULT_CONFIG.sorting,
            sorters=sorters,
            press_machines=press_machines,
            pallet_buffer_capacity=pallet_buffer_capacity,
        )
        melting = dataclasses.replace(
            DEFAULT_CONFIG.melting,
            furnace_count=furnace_count,
            batch_ton=batch_ton,
            pallets_per_batch=int(batch_ton / DEFAULT_CONFIG.sorting.pallet_ton),
        )
        casting = dataclasses.replace(
            DEFAULT_CONFIG.casting,
            flake_ratio=flake_ratio / 100.0,
            scr_ratio=1.0 - flake_ratio / 100.0,
        )
        outbound = dataclasses.replace(
            DEFAULT_CONFIG.outbound,
            empty_truck_interval_min=float(empty_truck_interval),
        )
        cfg = SimulationConfig(
            sim_days=sim_days,
            random_seed=int(random_seed),
            inbound=inbound,
            sorting=sorting,
            melting=melting,
            casting=casting,
            outbound=outbound,
        )

        with st.status(
            f"🔄 {sim_days}일치 시뮬레이션 실행 중",
            expanded=True,
        ) as sim_status:
            sim_status.write(
                "실제 벽시계는 보통 수 초 이내입니다. 가상 시간은 설정한 일수만큼 "
                "한 번에 재생됩니다."
            )
            t0 = time.perf_counter()
            metrics = run_simulation(cfg, progress=sim_status.write)
            elapsed = time.perf_counter() - t0
            sim_status.update(
                label=f"✅ 시뮬레이션 완료 (약 {elapsed:.2f}초)",
                state="complete",
                expanded=False,
            )

        st.success(
            f"✅ 결과 반영 완료 — **시뮬레이션** 탭에서 KPI·표를, **시각자료** 탭에서 그래프를 확인하세요. "
            f"(실행 {elapsed:.2f}초)"
        )

        analysis = analyze(metrics, cfg)
        st.session_state[_GUNSAN_LAST_RUN_BUNDLE_KEY] = {
            "metrics": metrics,
            "analysis": analysis,
            "cfg": cfg,
        }

    _gunsan_bundle = st.session_state.get(_GUNSAN_LAST_RUN_BUNDLE_KEY)
    _gunsan_show_results = _gunsan_bundle is not None

    tab_sim, tab_viz, tab_glossary = st.tabs(["시뮬레이션", "시각자료", "용어약어"])
    with tab_sim:
        if not _gunsan_show_results:
            _render_simulation_prerun_tabs(
                trucks_per_day=trucks_per_day,
                payload_ton=payload_ton,
                flake_ratio=flake_ratio,
                furnace_count=furnace_count,
            )
        if _gunsan_show_results:
            metrics = _gunsan_bundle["metrics"]
            analysis = _gunsan_bundle["analysis"]
            cfg = _gunsan_bundle["cfg"]
            sim_days = cfg.sim_days
            random_seed = cfg.random_seed
            trucks_per_day = cfg.inbound.trucks_per_day
            payload_ton = cfg.inbound.payload_ton
            unloading_bays = cfg.inbound.unloading_bays
            sorters = cfg.sorting.sorters
            press_machines = cfg.sorting.press_machines
            pallet_buffer_capacity = cfg.sorting.pallet_buffer_capacity
            furnace_count = cfg.melting.furnace_count
            batch_ton = cfg.melting.batch_ton
            flake_ratio = int(round(cfg.casting.flake_ratio * 100))
            empty_truck_interval = int(round(cfg.outbound.empty_truck_interval_min))
            summary = analysis.summary

            with st.expander("📌 이 시뮬 숫자가 어디서 오나요? (쉬운 설명)", expanded=False):
                _render_simulation_inputs_layperson_guide()
            _render_process_flow_section_header()
            _render_process_flow_pipeline_html(
                trucks_per_day=trucks_per_day,
                payload_ton=payload_ton,
                flake_ratio=flake_ratio,
                furnace_count=furnace_count,
                bottleneck=None,
                show_tooltips=True,
            )

            # ===== KPI 카드 (4열×3행, 10개 지표) =====
            st.header("📊 핵심 지표")
            _kpi_specs: list[tuple[str, str, str | None]] = [
                (
                    "처리 트럭 (입고)",
                    f"{summary['trucks_in_processed']} 대",
                    "시뮬레이션 기간 동안 입고 후 출차 완료된 트럭의 총 대수입니다. 1차 계근 → 하역 → 2차 계근 → 출차 과정을 완료한 트럭만 집계됩니다.",
                ),
                (
                    "출하 트럭",
                    f"{summary['trucks_out_dispatched']} 대",
                    "완제품을 적재하고 출고된 출하 트럭의 총 대수입니다. 빈 트럭 도착 → 상차 → 계근 → 출고 과정을 완료한 트럭만 집계됩니다.",
                ),
                (
                    "완료 배치",
                    f"{summary['melt_batches_completed']} 회",
                    "반사로에서 완료된 용해 배치의 총 횟수입니다. 1배치 = 장입 → 용해·정련 병목(약 13시간) → 주조 완료. 생산 능력의 핵심 지표입니다.",
                ),
                (
                    "총 생산량",
                    f"{summary['total_product_ton']:.0f} t",
                    "큐프레이크와 SCR 코일을 합한 총 생산량(톤)입니다. 출하된 제품 + 야적장 재고를 포함합니다.",
                ),
                (
                    "일평균 처리량",
                    f"{summary['throughput_ton_per_day']:.1f} t/일",
                    "일평균 생산량 = 총 생산량 ÷ 시뮬레이션 일수. 공장의 실질적인 생산 능력을 나타내는 핵심 KPI입니다.",
                ),
                (
                    "큐프레이크",
                    f"{summary['flake_ton']:.0f} t",
                    "생산된 Cu 플레이크의 총 중량(톤)입니다. 25kg 포대 단위로 생산되며, 전해 정련 및 합금 제조에 사용됩니다.",
                ),
                (
                    "SCR 코일",
                    f"{summary['scr_ton']:.0f} t",
                    "생산된 SCR(South Wire Rod) 코일의 총 중량(톤)입니다. 4톤/코일 단위로 생산되며, 전선/케이블 제조에 사용됩니다.",
                ),
                (
                    "입고 평균체류",
                    _format_kpi_minutes(float(summary["avg_truck_in_lead_min"])),
                    "입고 트럭의 평균 체류시간(도착~출차). 대기시간이 길면 하역장/계근대 병목을 의심해야 합니다. 목표: 60분 이내.",
                ),
                (
                    "출하 평균체류",
                    _format_kpi_minutes(float(summary["avg_truck_out_lead_min"])),
                    "출하 트럭의 평균 체류시간(도착~출고). 상차 대기, 제품 부족 등으로 지연될 수 있습니다. 목표: 90분 이내.",
                ),
                (
                    "배치 평균시간",
                    _format_kpi_minutes(float(summary["avg_melt_batch_min"])),
                    "반사로 1배치 완료에 걸리는 평균 시간(분). 장입(2~3h) + 용해(12h) + 주조(8h) ≈ 22~24시간이 정상입니다.",
                ),
            ]
            for _row in range(3):
                _kcols = st.columns(4)
                for _j in range(4):
                    _idx = _row * 4 + _j
                    if _idx >= len(_kpi_specs):
                        break
                    _label, _value, _help = _kpi_specs[_idx]
                    with _kcols[_j]:
                        st.metric(_label, _value, help=_help)

            st.header("📌 시뮬레이션 분석 결과 인사이트")
            st.caption(
                "KPI·자원 가동률·버퍼·일별 생산·트럭 이벤트를 바탕으로 규칙을 적용해 자동 생성한 요약입니다. "
                "투자·인력·일정 결정은 반드시 현장 데이터와 함께 검토하세요."
            )
            _ins_col, _rec_col = st.columns(2)
            with _ins_col:
                st.subheader("관찰 포인트")
                if analysis.insights:
                    for _txt in analysis.insights:
                        st.info(_txt)
                else:
                    st.caption("이번 실행에서 규칙으로 잡힌 특이 관찰이 없습니다.")
            with _rec_col:
                st.subheader("권장 액션")
                if analysis.recommendations:
                    for _txt in analysis.recommendations:
                        st.warning(_txt)
                else:
                    st.caption("이번 실행에서 자동 권장 문구가 없습니다.")
            with st.expander("🚛 트럭·물류 시사점 (누적 도착·출차·체류)", expanded=False):
                if analysis.truck_flow_insights:
                    for _txt in analysis.truck_flow_insights:
                        st.markdown(f"- {_txt}")
                else:
                    st.caption("트럭 흐름용 해설이 없습니다.")

            st.header("📖 일반인을 위한 상세 결과 해석")
            st.caption(
                "이번 실행의 지표를 처음 보는 분도 흐름을 따라갈 수 있도록 풀어 썼습니다. "
                "가동률·버퍼·트럭 등 **그래프는 옆의 시각자료 탭**에서 보시고, 아래 문단에서는 같은 내용을 말로 풀어 봅니다. "
                "HTML 리포트의 ‘일반인’ 절에도 같은 그림이 붙어 있습니다. "
                "투자·안전·계약 등 중요한 결정에는 참고용으로만 쓰고, 반드시 현장 검토가 필요합니다."
            )
            st.markdown("##### 이번 실행 그래프")
            st.info(
                "Plotly 그래프(가동률, 버퍼, 누적 트럭, 체류 분포, 일별 생산, 반사로 간트)는 "
                "**시각자료** 탭에서 한곳에 모아 보실 수 있습니다."
            )
            st.markdown(
                layperson_interpretation_markdown(metrics, cfg, analysis),
                unsafe_allow_html=True,
            )

            st.header("📜 시간순 사건 로그")
            st.caption(
                "시뮬레이션 동안 공장 안에서 일어난 일을 시간 순서대로 정리한 표입니다. "
                "공정 구간이나 키워드로 걸러 보거나 CSV·Markdown 보고서로 내려받을 수 있습니다."
            )
            _ev_df_all = _build_events_log_df(metrics)
            _report_ev_df = _ev_df_all.iloc[0:0]
            _report_ev_total = int(len(_ev_df_all))
            _report_ev_note = "이번 실행에서 기록된 사건이 없습니다."
            _display_cols = ["#", "시각", "경과(분)", "구간", "사건", "상세 설명"]
            if _ev_df_all.empty:
                st.info("이번 실행에서 기록된 사건이 없습니다.")
            else:
                _stages = sorted(_ev_df_all["구간"].unique().tolist())
                _f1, _f2 = st.columns(2)
                with _f1:
                    _stage_pick = st.multiselect(
                        "공정 구간으로 보기",
                        options=_stages,
                        default=_stages,
                        key="gunsan_event_log_stages",
                        help="보고 싶은 공정 단계만 골라 표시합니다. 비워 두면 표가 비어 있게 됩니다.",
                    )
                with _f2:
                    _needle = st.text_input(
                        "키워드 검색 (사건·상세 설명에서 부분 일치)",
                        "",
                        key="gunsan_event_log_search",
                        placeholder="예) 12번 트럭, 1호 반사로, 후레이크, 버퍼",
                        help="입력한 단어가 사건 이름이나 상세 설명에 포함된 행만 보여 줍니다.",
                    )
                if not _stage_pick:
                    st.warning("최소 한 개 이상의 공정 구간을 선택해 주세요.")
                    _df_view = _ev_df_all.iloc[0:0]
                    _report_ev_note = "공정 구간이 선택되지 않아 표가 비어 있습니다."
                else:
                    _df_view = _ev_df_all[_ev_df_all["구간"].isin(_stage_pick)]
                    if _needle.strip():
                        _q = _needle.strip().casefold()
                        _mask = (
                            _df_view["사건"].astype(str).str.casefold().str.contains(_q, na=False)
                            | _df_view["상세 설명"].astype(str).str.casefold().str.contains(_q, na=False)
                        )
                        _df_view = _df_view[_mask]
                    _bits = [f"선택 구간: {', '.join(_stage_pick)}"]
                    if _needle.strip():
                        _bits.append(f"키워드 검색: `{_needle.strip()}`")
                    _report_ev_note = " · ".join(_bits)
                _report_ev_df = _df_view
                st.caption(f"표시 {len(_df_view):,}건 / 전체 {len(_ev_df_all):,}건")
                # 화면에는 사람이 읽기 쉬운 한글 컬럼만 보이고, 영문 원본은 CSV/분석용으로 숨겨 둔다.
                st.dataframe(
                    _df_view[_display_cols],
                    use_container_width=True,
                    height=420,
                    hide_index=True,
                )
                with st.expander("표에서 쓰인 용어 풀이", expanded=False):
                    st.markdown(
                        "- 시각: 시뮬 시작을 0일차 00:00 으로 본 시계. 예) `3일차 14:30`\n"
                        "- 경과(분): 시뮬 시작부터 흐른 시간(분)\n"
                        "- 구간: 공장 안에서 사건이 일어난 공정 단계\n"
                        "- 사건: 그 구간에서 일어난 일\n"
                        "- 상세 설명: 트럭 번호·반사로 번호·생산량·버퍼 상태 등 함께 기록된 정보"
                    )

            _ev_md_slice = (
                _report_ev_df[_display_cols]
                if not _report_ev_df.empty
                and all(c in _report_ev_df.columns for c in _display_cols)
                else _report_ev_df
            )
            _report_now = datetime.now()
            _report_fn_ts = _report_now.strftime("%Y%m%d_%H%M%S")
            _report_gen_stamp = _report_now.strftime("%Y-%m-%d %H:%M:%S")
            _md_report = _build_simulation_results_markdown(
                metrics,
                cfg,
                analysis,
                _ev_md_slice,
                _report_ev_total,
                _report_ev_note,
                _kpi_specs,
                generated_at=_report_gen_stamp,
            )

            _dl_csv, _dl_md, _dl_pdf = st.columns(3)
            with _dl_csv:
                if not _ev_df_all.empty:
                    _csv_bytes = _report_ev_df.to_csv(index=False).encode("utf-8-sig")
                    st.download_button(
                        label="CSV 다운로드 (현재 필터 적용)",
                        data=_csv_bytes,
                        file_name="gunsan_sim_event_log.csv",
                        mime="text/csv",
                        key="gunsan_event_log_csv",
                        help="현재 필터·검색 조건에 맞는 행만 CSV로 내려받습니다. 영문 원본 stage/kind 컬럼도 함께 포함됩니다.",
                    )
                else:
                    st.caption("사건 로그가 없어 CSV에 담을 행이 없습니다.")
            with _dl_md:
                st.download_button(
                    label="Markdown 보고서",
                    data=_md_report.encode("utf-8"),
                    file_name=f"gunsan_sim_report_{_report_fn_ts}.md",
                    mime="text/markdown",
                    key="gunsan_md_full_report",
                    help="KPI·해석·병목·표·인사이트·(현재 필터와 동일한) 사건 로그. Plotly 차트는 제외.",
                )
            with _dl_pdf:
                _pdf_bytes, _pdf_err = markdown_simulation_report_to_pdf(_md_report)
                if _pdf_bytes:
                    st.download_button(
                        label="PDF 보고서",
                        data=_pdf_bytes,
                        file_name=f"gunsan_sim_report_{_report_fn_ts}.pdf",
                        mime="application/pdf",
                        key="gunsan_pdf_full_report",
                        help="Markdown 보고서와 동일 내용을 A4 PDF로 저장합니다. 한글 글꼴(맑은 고딕 등)이 필요합니다.",
                    )
                else:
                    st.caption(_pdf_err or "PDF를 생성할 수 없습니다.")

            # ===== 병목 진단 =====
            st.header("🔍 병목 진단")
            with st.expander("ℹ️ 병목(Bottleneck)이란?", expanded=False):
                st.markdown("""
                병목은 전체 공정의 처리량을 제한하는 가장 느린 공정 단계입니다.
        
                - 병목 자원의 가동률이 가장 높음 (90% 이상)
                - 병목 앞단에 대기열/재고가 누적됨
                - 병목 개선이 전체 처리량 향상에 직결됨
        
                일반적인 병목 순서: 반사로 용해 > 압착기 > 하역장 > 선별
                """)
            st.error(f"식별된 병목: {analysis.bottleneck} — {analysis.bottleneck_reason}")
            st.caption(
                "아래 HTML 공정 파이프라인 카드에서 병목 단계가 강조됩니다. "
                "페이지 상단 「세부공정 프로세스」 차트는 흐름만 보여 주며 병목 색은 적용되지 않습니다."
            )
            _render_process_flow_pipeline_html(
                trucks_per_day=trucks_per_day,
                payload_ton=payload_ton,
                flake_ratio=flake_ratio,
                furnace_count=furnace_count,
                bottleneck=analysis.bottleneck,
                show_tooltips=True,
            )

    with tab_viz:
        if not _gunsan_show_results:
            st.info(
                "아직 실행된 시뮬레이션 결과가 없습니다. "
                "왼쪽 사이드바에서 **시뮬레이션 실행**을 누른 뒤, 이 **시각자료** 탭으로 돌아오면 그래프가 표시됩니다."
            )
        else:
            _vb = st.session_state[_GUNSAN_LAST_RUN_BUNDLE_KEY]
            st.header("📈 시각자료 (이번 실행)")
            st.caption(
                "막대는 자원 가동률, 곡선은 버퍼·누적 트럭, 히스토그램은 트럭 체류시간, "
                "막대 누적은 일별 생산, 가로 막대는 반사로 일정입니다. 범례를 눌러 선을 끄거나 호버로 숫자를 확인할 수 있습니다."
            )
            _render_layperson_plotly_figures(
                _vb["metrics"],
                _vb["cfg"],
                _vb["analysis"],
                chart_key_prefix="gunsan_viz_tab",
            )

    with tab_glossary:
        _render_glossary_page()
